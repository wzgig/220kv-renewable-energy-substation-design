from __future__ import annotations

import re
import unittest
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

from report.scripts.build_reports import add_bullets


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
BASE_TITLE = "220kV新能源汇集变电所电气一次部分初步设计"
DEFENSE_TITLE = f"{BASE_TITLE}答辩汇报"
PUBLIC_AUTHOR = "课程设计项目组（公开版）"

PDF_PAGES = {
    "01_220kV新能源汇集变电所_技术设计说明书.pdf": 35,
    "02_220kV新能源汇集变电所_技术设计计算书.pdf": 29,
    "03_220kV新能源汇集变电所_课程设计总结.pdf": 4,
    "04_220kV新能源汇集变电所_答辩汇报.pdf": 11,
    "05_220kV新能源汇集变电所_答辩问题清单.pdf": 7,
}

DOCX_FILES = (
    "01_220kV新能源汇集变电所_技术设计说明书.docx",
    "02_220kV新能源汇集变电所_技术设计计算书.docx",
    "03_220kV新能源汇集变电所_课程设计总结.docx",
    "05_220kV新能源汇集变电所_答辩问题清单.docx",
)

PPTX_FILE = "04_220kV新能源汇集变电所_答辩汇报.pptx"


def core_properties(archive: ZipFile) -> dict[str, str]:
    root = ET.fromstring(archive.read("docProps/core.xml"))
    return {element.tag.rsplit("}", 1)[-1]: element.text or "" for element in root}


def document_text(path: Path) -> str:
    with ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    return "".join(
        element.text or ""
        for element in root.iter()
        if element.tag.rsplit("}", 1)[-1] == "t"
    )


class ReportDeliverableTests(unittest.TestCase):
    def test_expected_report_files_and_reproduction_scripts_exist(self) -> None:
        expected = set(PDF_PAGES) | set(DOCX_FILES) | {PPTX_FILE}
        for name in expected:
            path = REPORT_DIR / name
            self.assertTrue(path.is_file(), name)
            self.assertGreater(path.stat().st_size, 1024, name)

        for name in (
            "build_reports.py",
            "export_reports.ps1",
            "sanitize_public_metadata.py",
        ):
            self.assertTrue((REPORT_DIR / "scripts" / name).is_file(), name)

    def test_pdf_page_counts_and_public_metadata(self) -> None:
        for name, expected_pages in PDF_PAGES.items():
            with self.subTest(name=name):
                reader = PdfReader(REPORT_DIR / name)
                metadata = dict(reader.metadata or {})
                self.assertEqual(len(reader.pages), expected_pages)
                self.assertEqual(metadata.get("/Author"), PUBLIC_AUTHOR)
                expected_title = DEFENSE_TITLE if name.startswith("04_") else BASE_TITLE
                self.assertEqual(metadata.get("/Title"), expected_title)
                self.assertNotEqual(metadata.get("/Title"), "Presentation")

    def test_docx_packages_have_public_privacy_metadata(self) -> None:
        rsid_attribute = re.compile(rb"\s+w:rsid[A-Za-z0-9]*=[\"'][^\"']*[\"']")
        for name in DOCX_FILES:
            with self.subTest(name=name), ZipFile(REPORT_DIR / name) as archive:
                self.assertIsNone(archive.testzip())
                self.assertNotIn("docProps/custom.xml", archive.namelist())
                properties = core_properties(archive)
                self.assertEqual(properties.get("creator", ""), "")
                self.assertEqual(properties.get("lastModifiedBy", ""), "")

                app_root = ET.fromstring(archive.read("docProps/app.xml"))
                app_properties = {
                    element.tag.rsplit("}", 1)[-1]: element.text or ""
                    for element in app_root
                }
                pdf_name = Path(name).with_suffix(".pdf").name
                self.assertEqual(
                    int(app_properties.get("Pages", "0")), PDF_PAGES[pdf_name]
                )

                for member in archive.namelist():
                    if member.startswith("word/") and member.endswith(".xml"):
                        self.assertIsNone(rsid_attribute.search(archive.read(member)), member)

    def test_report_content_contains_latest_grounding_and_duty_boundaries(self) -> None:
        payload = "\n".join(document_text(REPORT_DIR / name) for name in DOCX_FILES)
        for expected in (
            "1000kVA ZN接地变",
            "ZCT",
            "母联合闸前",
            "7.385",
            "18.798",
            "2230.732",
            "2404.371",
            "2225.155",
            "2398.360",
            "2221.564",
            "2394.612",
            "2342.268",
            "2524.590",
            "3117.691",
            "机械/电动力待核",
        ):
            self.assertIn(expected, payload)

    def test_design_and_calculation_books_have_formal_narrative_depth(self) -> None:
        description = Document(REPORT_DIR / DOCX_FILES[0])
        calculation = Document(REPORT_DIR / DOCX_FILES[1])

        def metrics(doc: Document) -> tuple[int, int, int, list[str], str]:
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            body_text = "\n".join(
                p.text.strip()
                for p in paragraphs
                if not p.style.name.startswith(("Heading", "Caption", "toc"))
            )
            numbered_equations = sum(
                1 for p in paragraphs if re.search(r"\(\d+-\d+\)\s*$", p.text.strip())
            )
            headings = [
                p.text.strip() for p in paragraphs if p.style.name.startswith("Heading")
            ]
            return len(body_text), numbered_equations, len(doc.tables), headings, body_text

        desc_chars, desc_equations, desc_tables, desc_headings, desc_text = metrics(
            description
        )
        calc_chars, calc_equations, calc_tables, calc_headings, calc_text = metrics(
            calculation
        )

        self.assertGreaterEqual(desc_chars, 12000)
        self.assertGreaterEqual(desc_equations, 25)
        self.assertGreaterEqual(desc_tables, 25)
        self.assertIn("结果分析：", desc_text)
        self.assertIn("第六章 结论与设计边界", desc_headings)

        self.assertGreaterEqual(calc_chars, 13000)
        self.assertGreaterEqual(calc_equations, 50)
        self.assertGreaterEqual(calc_tables, 18)
        self.assertIn("结果分析：", calc_text)
        self.assertIn("第一章 负荷及变压器容量计算", calc_headings)
        self.assertIn("第四章 主要电气设备选择与校验", calc_headings)
        self.assertNotIn("前置计算", calc_text)

    def test_front_matter_and_main_body_use_distinct_page_number_formats(self) -> None:
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for name in DOCX_FILES[:2]:
            with self.subTest(name=name), ZipFile(REPORT_DIR / name) as archive:
                root = ET.fromstring(archive.read("word/document.xml"))
                page_number_types = []
                for section in root.findall(".//w:sectPr", namespace):
                    node = section.find("w:pgNumType", namespace)
                    if node is None:
                        continue
                    page_number_types.append(
                        (
                            node.get(f"{{{namespace['w']}}}fmt"),
                            node.get(f"{{{namespace['w']}}}start"),
                        )
                    )

                self.assertIn(("lowerRoman", "1"), page_number_types)
                self.assertTrue(
                    any(fmt in (None, "decimal") and start == "1" for fmt, start in page_number_types)
                )

                footer_payload = b"\n".join(
                    archive.read(member)
                    for member in archive.namelist()
                    if member.startswith("word/footer") and member.endswith(".xml")
                )
                self.assertGreaterEqual(footer_payload.count(b" PAGE "), 2)

    def test_ordered_list_blocks_restart_from_one(self) -> None:
        doc = Document()
        add_bullets(doc, ["步骤甲", "步骤乙"], ordered=True)
        add_bullets(doc, ["步骤丙", "步骤丁"], ordered=True)

        num_ids = []
        for paragraph in doc.paragraphs:
            num_pr = paragraph._p.get_or_add_pPr().find(qn("w:numPr"))
            self.assertIsNotNone(num_pr)
            num_id = num_pr.find(qn("w:numId"))
            self.assertIsNotNone(num_id)
            num_ids.append(num_id.get(qn("w:val")))

        self.assertEqual(num_ids[0], num_ids[1])
        self.assertEqual(num_ids[2], num_ids[3])
        self.assertNotEqual(num_ids[0], num_ids[2])

        numbering = doc.part.numbering_part.element
        for num_id in {num_ids[0], num_ids[2]}:
            num = next(
                node
                for node in numbering.findall(qn("w:num"))
                if node.get(qn("w:numId")) == num_id
            )
            abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
            abstract = next(
                node
                for node in numbering.findall(qn("w:abstractNum"))
                if node.get(qn("w:abstractNumId")) == abstract_id
            )
            start = abstract.find(qn("w:lvl")).find(qn("w:start"))
            self.assertEqual(start.get(qn("w:val")), "1")
            start_override = num.find(qn("w:lvlOverride")).find(
                qn("w:startOverride")
            )
            self.assertEqual(start_override.get(qn("w:val")), "1")

    def test_final_description_preserves_list_numbering_semantics(self) -> None:
        path = REPORT_DIR / DOCX_FILES[0]
        description = Document(path)

        def paragraph_num_id(prefix: str) -> str:
            paragraph = next(
                p for p in description.paragraphs if p.text.strip().startswith(prefix)
            )
            num_pr = paragraph._p.find("./" + qn("w:pPr") + "/" + qn("w:numPr"))
            self.assertIsNotNone(num_pr)
            num_id = num_pr.find(qn("w:numId"))
            self.assertIsNotNone(num_id)
            return num_id.get(qn("w:val"))

        conclusion_num_id = paragraph_num_id("主接线采用220kV")
        accident_num_id = paragraph_num_id("保护动作切除故障设备")
        self.assertNotEqual(conclusion_num_id, accident_num_id)

        numbering = description.part.numbering_part.element

        def numbering_definition(num_id: str):
            num = next(
                node
                for node in numbering.findall(qn("w:num"))
                if node.get(qn("w:numId")) == num_id
            )
            abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
            abstract = next(
                node
                for node in numbering.findall(qn("w:abstractNum"))
                if node.get(qn("w:abstractNumId")) == abstract_id
            )
            return num, abstract

        _, conclusion_abstract = numbering_definition(conclusion_num_id)
        accident_num, accident_abstract = numbering_definition(accident_num_id)
        self.assertEqual(
            conclusion_abstract.find(qn("w:lvl"))
            .find(qn("w:numFmt"))
            .get(qn("w:val")),
            "bullet",
        )
        self.assertEqual(
            accident_abstract.find(qn("w:lvl"))
            .find(qn("w:numFmt"))
            .get(qn("w:val")),
            "decimal",
        )
        start_override = accident_num.find(qn("w:lvlOverride")).find(
            qn("w:startOverride")
        )
        self.assertEqual(start_override.get(qn("w:val")), "1")

        with ZipFile(path) as archive:
            numbering_root = ET.fromstring(archive.read("word/numbering.xml"))
        child_tags = [element.tag.rsplit("}", 1)[-1] for element in numbering_root]
        first_num_index = child_tags.index("num")
        last_abstract_index = max(
            index for index, tag in enumerate(child_tags) if tag == "abstractNum"
        )
        self.assertLess(last_abstract_index, first_num_index)

    def test_design_description_layout_controls_are_persisted(self) -> None:
        description = Document(REPORT_DIR / DOCX_FILES[0])
        headings = {p.text.strip(): p for p in description.paragraphs}
        for heading_text in (
            "1.5 正常、条件与禁止运行方式",
            "4.5 母线、绝缘子与套管",
        ):
            page_break = headings[heading_text]._p.get_or_add_pPr().find(
                qn("w:pageBreakBefore")
            )
            if page_break is not None:
                self.assertEqual(page_break.get(qn("w:val")), "0")

        switchgear_table = next(
            table
            for table in description.tables
            if table.rows[0].cells[0].text.strip() == "安装位置"
            and table.rows[0].cells[-1].text.strip() == "热稳定"
        )
        for row in switchgear_table.rows[:2]:
            for column in (2, 3, 4, 5):
                self.assertIsNotNone(row.cells[column]._tc.tcPr.find(qn("w:noWrap")))

    def test_pptx_is_valid_public_and_has_no_generator_theme_name(self) -> None:
        path = REPORT_DIR / PPTX_FILE
        with ZipFile(path) as archive:
            self.assertIsNone(archive.testzip())
            properties = core_properties(archive)
            self.assertEqual(properties.get("title"), DEFENSE_TITLE)
            self.assertEqual(properties.get("creator"), PUBLIC_AUTHOR)
            self.assertEqual(properties.get("lastModifiedBy"), PUBLIC_AUTHOR)

            slide_count = sum(
                1
                for name in archive.namelist()
                if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
            )
            self.assertEqual(slide_count, 11)

            expected_theme_name = "课程设计主题"
            for name in archive.namelist():
                if not name.startswith("ppt/theme/") or not name.endswith(".xml"):
                    continue
                theme_root = ET.fromstring(archive.read(name))
                for element in theme_root.iter():
                    if element.tag.rsplit("}", 1)[-1] in {"theme", "clrScheme", "fmtScheme"}:
                        self.assertEqual(element.attrib.get("name"), expected_theme_name)

            self.assertIn(expected_theme_name.encode("utf-8"), archive.read("docProps/app.xml"))

            core_xml = archive.read("docProps/core.xml")
            self.assertIn(b"xmlns:dcterms=", core_xml)
            self.assertIn(b'xsi:type="dcterms:W3CDTF"', core_xml)


if __name__ == "__main__":
    unittest.main()
