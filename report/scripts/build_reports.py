from __future__ import annotations

import json
import math
import re
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import yaml
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "materials-private" / "02_templates" / "课设模板.docx"
REPORT_DIR = ROOT / "report"
LOAD_JSON = ROOT / "calculations" / "results" / "load_and_transformer_results.json"
SC_JSON = ROOT / "calculations" / "results" / "short_circuit" / "short_circuit_results.json"
EQUIP_JSON = (
    ROOT
    / "calculations"
    / "results"
    / "equipment_selection"
    / "equipment_selection_results.json"
)
BASELINE_YAML = ROOT / "data" / "design_baseline.yaml"
INPUTS_YAML = ROOT / "data" / "design_inputs.yaml"

TITLE = "220kV新能源汇集变电所电气一次部分初步设计"
SHORT_TITLE = "220kV新能源汇集变电所课程设计"
DATE_RANGE = "2026年7月6日—2026年7月17日"
PUBLIC_NOTE = "公开版：姓名、学号、班级及签字字段留空，提交时由学生手工填写。"
ENGINEERING_NOTE = "课程设计底稿，非施工图；工程实施前须按现行标准及厂家资料复核。"


def load_data() -> dict:
    return {
        "load": json.loads(LOAD_JSON.read_text(encoding="utf-8")),
        "sc": json.loads(SC_JSON.read_text(encoding="utf-8")),
        "equipment": json.loads(EQUIP_JSON.read_text(encoding="utf-8")),
        "baseline": yaml.safe_load(BASELINE_YAML.read_text(encoding="utf-8")),
        "inputs": yaml.safe_load(INPUTS_YAML.read_text(encoding="utf-8")),
    }


def fmt(value: float | int | None, digits: int = 3) -> str:
    if value is None:
        return "待定"
    if isinstance(value, int):
        return str(value)
    return f"{value:.{digits}f}"


def remove_all_body_content(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)
    no_wrap.set(qn("w:val"), "true")


def set_inline_shape_alt_text(inline_shape, alt_text: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Cm(width_cm)


def set_table_width(table, widths_cm: list[float]) -> None:
    total_dxa = sum(int(Cm(w).emu / 635) for w in widths_cm)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(Cm(width).emu / 635)))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_cm[index])


def add_bottom_border(paragraph, color="000000", size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_east_asia_font(run, name="宋体", latin="Times New Roman", size=12, bold=None):
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:cs"), latin)
    return run


def style_document(doc: Document) -> None:
    settings = doc.settings._element
    mirror_margins = settings.find(qn("w:mirrorMargins"))
    if mirror_margins is not None:
        settings.remove(mirror_margins)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.74)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    for name, size, align, before, after in (
        ("Heading 1", 16, WD_ALIGN_PARAGRAPH.CENTER, 12, 8),
        ("Heading 2", 14, WD_ALIGN_PARAGRAPH.LEFT, 10, 6),
        ("Heading 3", 12, WD_ALIGN_PARAGRAPH.LEFT, 8, 4),
    ):
        st = styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.paragraph_format.alignment = align
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.left_indent = Cm(0)
        st.paragraph_format.right_indent = Cm(0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True
    styles["Heading 1"].paragraph_format.page_break_before = None

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10.5)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = True

    for name, size, left in (("toc 1", 10.5, 0.0), ("toc 2", 10.0, 0.65), ("toc 3", 9.5, 1.3)):
        try:
            toc_style = styles[name]
        except KeyError:
            continue
        toc_style.font.name = "Times New Roman"
        toc_style.font.size = Pt(size)
        toc_style.font.color.rgb = RGBColor(0, 0, 0)
        toc_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        toc_style.paragraph_format.left_indent = Cm(left)
        toc_style.paragraph_format.first_line_indent = Cm(0)
        toc_style.paragraph_format.space_before = Pt(0)
        toc_style.paragraph_format.space_after = Pt(0)
        toc_style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        set_section_geometry(section)


def set_section_geometry(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def set_page_number_start(section, value: int = 1, fmt: str = "decimal") -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(value))
    pg_num_type.set(qn("w:fmt"), fmt)


def add_update_fields_setting(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def configure_cover_section(section) -> None:
    set_section_geometry(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.clear()
    for p in section.footer.paragraphs:
        p.clear()


def configure_front_section(section) -> None:
    configure_cover_section(section)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)
    for run in fp.runs:
        set_east_asia_font(run, name="宋体", size=9)
    set_page_number_start(section, 1, fmt="lowerRoman")


def configure_main_section(section, running_title: str) -> None:
    set_section_geometry(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_east_asia_font(hp.add_run(running_title), name="宋体", size=9)
    add_bottom_border(hp, size="6")
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)
    for run in fp.runs:
        set_east_asia_font(run, name="宋体", size=9)
    set_page_number_start(section, 1, fmt="decimal")


def compact_break_paragraph(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_pr = paragraph._p.get_or_add_pPr()
    snap_to_grid = p_pr.find(qn("w:snapToGrid"))
    if snap_to_grid is None:
        snap_to_grid = OxmlElement("w:snapToGrid")
        p_pr.append(snap_to_grid)
    snap_to_grid.set(qn("w:val"), "0")
    widow_control = p_pr.find(qn("w:widowControl"))
    if widow_control is None:
        widow_control = OxmlElement("w:widowControl")
        p_pr.append(widow_control)
    widow_control.set(qn("w:val"), "0")


def start_main_section(doc: Document, running_title: str):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    compact_break_paragraph(doc.paragraphs[-1])
    configure_main_section(section, running_title)
    return section


def add_toc(doc: Document, levels="1-3") -> None:
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = Cm(0)
    heading.paragraph_format.space_after = Pt(12)
    set_east_asia_font(heading.add_run("目 录"), name="黑体", size=16, bold=True)
    p = doc.add_paragraph()
    compact_break_paragraph(p)
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在 Microsoft Word 中自动更新。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, placeholder, end])


def logo_bytes() -> tuple[bytes, bytes]:
    with ZipFile(TEMPLATE) as archive:
        return archive.read("word/media/image1.png"), archive.read("word/media/image2.png")


def add_cover(doc: Document, document_type: str) -> None:
    logo, calligraphy = logo_bytes()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    logo_shape = p.add_run().add_picture(BytesIO(logo), width=Cm(2.6))
    set_inline_shape_alt_text(logo_shape, "学校校徽")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    calligraphy_shape = p.add_run().add_picture(BytesIO(calligraphy), width=Cm(8.5))
    set_inline_shape_alt_text(calligraphy_shape, "学校名称书法标识")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_east_asia_font(p.add_run("《发电厂电气部分》课程设计"), name="黑体", size=22, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    set_east_asia_font(p.add_run(TITLE), name="黑体", size=18, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_east_asia_font(p.add_run(document_type), name="黑体", size=20, bold=True)

    labels = [
        ("学    院", "电气与信息工程学院"),
        ("专    业", "电气工程及其自动化"),
        ("班    级", ""),
        ("姓    名", ""),
        ("学    号", ""),
        ("同组设计者", ""),
        ("指导教师", ""),
        ("任务起止日期", DATE_RANGE),
    ]
    table = doc.add_table(rows=len(labels), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [4.0, 9.2])
    for row, (label, value) in zip(table.rows, labels):
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        p0.paragraph_format.first_line_indent = Cm(0)
        p0.paragraph_format.line_spacing = 1.0
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        set_east_asia_font(p0.add_run(label), name="宋体", size=12)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.first_line_indent = Cm(0)
        p1.paragraph_format.line_spacing = 1.0
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        set_east_asia_font(p1.add_run(value if value else "____________________________"), name="宋体", size=12)
        for cell in row.cells:
            set_cell_margins(cell, top=70, bottom=70)
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "nil")
                borders.append(el)
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single" if cell is row.cells[1] else "nil")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:color"), "000000")
            borders.append(bottom)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    set_east_asia_font(p.add_run(PUBLIC_NOTE), name="宋体", size=9)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_east_asia_font(p.add_run(ENGINEERING_NOTE), name="宋体", size=9)


def set_page_break_before(paragraph, enabled: bool) -> None:
    """Write an explicit OOXML value so Word cannot inherit the style setting."""
    p_pr = paragraph._p.get_or_add_pPr()
    page_break = p_pr.find(qn("w:pageBreakBefore"))
    if page_break is None:
        page_break = OxmlElement("w:pageBreakBefore")
        p_pr.append(page_break)
    page_break.set(qn("w:val"), "1" if enabled else "0")


def add_heading(
    doc: Document,
    text: str,
    level: int = 1,
    *,
    page_break_before: bool | None = None,
):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    if page_break_before is not None:
        set_page_break_before(paragraph, page_break_before)
    elif level == 1:
        set_page_break_before(paragraph, True)
    return paragraph


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_east_asia_font(p.add_run(bold_lead), name="宋体", size=12, bold=True)
        text = text[len(bold_lead) :]
    set_east_asia_font(p.add_run(text), name="宋体", size=12)
    return p


def get_list_num_id(doc: Document, ordered: bool) -> int:
    # Each ordered-list block gets its own numbering definition so it restarts
    # at 1. Bullet lists can safely share one definition throughout a document.
    cache_name = None if ordered else "_course_bullet_num_id"
    cached = getattr(doc, cache_name, None) if cache_name else None
    if cached is not None:
        return cached
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{0xC0DE0000 + abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    template_id = OxmlElement("w:tmpl")
    template_id.set(qn("w:val"), f"{0xB10C0000 + abstract_id:08X}")
    abstract.append(template_id)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "\uf0b7")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "420")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "300")
    p_pr.append(ind)
    lvl.append(p_pr)
    if not ordered:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Symbol")
        r_fonts.set(qn("w:hAnsi"), "Symbol")
        r_fonts.set(qn("w:hint"), "default")
        r_pr.append(r_fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    if ordered:
        lvl_override = OxmlElement("w:lvlOverride")
        lvl_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        lvl_override.append(start_override)
        num.append(lvl_override)
    numbering.append(num)
    if cache_name:
        setattr(doc, cache_name, num_id)
    return num_id


def add_bullets(
    doc: Document,
    items: list[str],
    ordered: bool = False,
    *,
    keep_together: bool = False,
) -> None:
    num_id = get_list_num_id(doc, ordered)
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.keep_together = keep_together
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num])
        p_pr.append(num_pr)
        set_east_asia_font(p.add_run(item), name="宋体", size=11.5)


def add_equation(
    doc: Document,
    expression: str,
    number: str | None = None,
    *,
    compact: bool = False,
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(1 if compact else 4)
    p.paragraph_format.space_after = Pt(1 if compact else 4)
    if number:
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(expression)
    set_east_asia_font(run, name="Cambria Math", latin="Cambria Math", size=11.0 if compact else 11.5)
    if number:
        p.add_run("\t")
        set_east_asia_font(p.add_run(f"({number})"), name="宋体", size=10.5)
    return p


def add_formula_terms(doc: Document, text: str):
    paragraph = add_body(doc, f"式中：{text}", bold_lead="式中：")
    paragraph.paragraph_format.keep_together = True
    return paragraph


def add_result_analysis(doc: Document, text: str):
    paragraph = add_body(doc, f"结果分析：{text}", bold_lead="结果分析：")
    paragraph.paragraph_format.keep_together = True
    return paragraph


def add_table(
    doc: Document,
    caption: str,
    headers: list[str],
    rows: list[list[object]],
    widths_cm: list[float] | None = None,
    font_size: float = 9.5,
    note: str | None = None,
    nowrap_header_columns: set[int] | None = None,
    nowrap_body_columns: set[int] | None = None,
    margin_x: int = 110,
    margin_y: int = 90,
    keep_note_with_table: bool = False,
) -> None:
    nowrap_header_columns = nowrap_header_columns or set()
    nowrap_body_columns = nowrap_body_columns or set()
    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    if widths_cm is None:
        widths_cm = [15.5 / len(headers)] * len(headers)
    if abs(sum(widths_cm) - 15.5) > 0.2:
        scale = 15.5 / sum(widths_cm)
        widths_cm = [w * scale for w in widths_cm]
    set_table_width(table, widths_cm)
    header = table.rows[0]
    set_repeat_table_header(header)
    set_row_cant_split(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        if index in nowrap_header_columns:
            set_cell_no_wrap(cell)
        set_cell_shading(cell, "D9E2F3")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        set_east_asia_font(p.add_run(str(value)), name="黑体", size=font_size, bold=True)
        set_cell_margins(cell, top=margin_y, bottom=margin_y, start=margin_x, end=margin_x)
    for row_values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            if index in nowrap_body_columns:
                set_cell_no_wrap(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_east_asia_font(p.add_run(str(value)), name="宋体", size=font_size)
            set_cell_margins(cell, top=margin_y, bottom=margin_y, start=margin_x, end=margin_x)
    if note:
        if keep_note_with_table:
            for cell in table.rows[-1].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.keep_together = True
        set_east_asia_font(p.add_run(f"注：{note}"), name="宋体", size=9)


def add_figure(doc: Document, path: Path, caption: str, width_cm: float = 15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    figure_shape = p.add_run().add_picture(str(path), width=Cm(width_cm))
    set_inline_shape_alt_text(figure_shape, caption)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = False
    cap.paragraph_format.keep_together = True


def add_optional_figure(
    doc: Document,
    path: Path,
    caption: str,
    *,
    missing_message: str,
    width_cm: float = 15.5,
) -> None:
    if path.exists():
        add_figure(doc, path, caption, width_cm=width_cm)
        return
    add_note_box(doc, "图纸待集成", missing_message)


def add_note_box(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [15.5])
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_east_asia_font(p.add_run(f"{title}："), name="黑体", size=10.5, bold=True)
    set_east_asia_font(p.add_run(text), name="宋体", size=10.5)


def add_references(
    doc: Document,
    references: list[str],
    *,
    page_break_before: bool = True,
) -> None:
    add_heading(doc, "参考文献", 1, page_break_before=page_break_before)
    for index, item in enumerate(references, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.keep_together = True
        set_east_asia_font(p.add_run(f"[{index}] {item}"), name="宋体", size=10.0)


REFERENCES = [
    "苗世洪等. 发电厂电气部分（第五版）[M]. 北京: 中国电力出版社, 2015.",
    "西北电力设计院. 电力工程电气设计手册: 电气一次部分[M]. 北京: 中国电力出版社.",
    "GB/T 15544.1-2023 三相交流系统短路电流计算 第1部分: 电流计算.",
    "DL/T 5222-2021 导体和电器选择设计规程.",
    "DL/T 5352-2018 高压配电装置设计规范.",
    "GB 50060-2008 3~110kV高压配电装置设计规范.",
    "GB/T 3906-2020 3.6kV~40.5kV交流金属封闭开关设备和控制设备.",
    "GB/T 1984-2024 高压交流断路器.",
    "GB/T 1985-2023 高压交流隔离开关和接地开关.",
    "GB/T 20840.1-2010 互感器 第1部分: 通用技术要求.",
    "GB/T 50064-2014 交流电气装置的过电压保护和绝缘配合设计规范.",
    "GB/T 50065-2011 交流电气装置的接地设计规范.",
]


def new_document(document_type: str) -> Document:
    doc = Document(TEMPLATE)
    remove_all_body_content(doc)
    style_document(doc)
    section = doc.sections[0]
    configure_cover_section(section)
    add_cover(doc, document_type)
    return doc


def build_design_description(data: dict) -> Path:
    load = data["load"]
    sc = data["sc"]
    equipment = data["equipment"]
    baseline = data["baseline"]
    load35 = load["load_35kv"]
    load10 = load["load_10kv"]
    main_transformer = load["main_transformer"]
    reactive = load["reactive_compensation"]
    station_service = load["station_service"]
    outgoing = load["outgoing_220kv"]
    outgoing_ampacity = load["outgoing_220kv_conductor_ampacity"]
    sc_points = sc["points"]
    course_completion = equipment["course_completion"]
    busbars = {item["id"]: item for item in course_completion["busbars"]}
    arresters = {item["id"]: item for item in course_completion["surge_arresters"]}
    grounding_packages = {
        item["id"]: item
        for item in course_completion["grounding_transformer_resistor_packages"]
    }
    doc = new_document("技术设计说明书")

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_front_section(front)
    add_heading(doc, "摘  要", 1, page_break_before=False)
    add_body(
        doc,
        "本课程设计面向一座220kV新能源汇集变电所，完成220/35/10/0.4kV电气一次部分初步设计。"
        "在任务书给定的系统容量、线路、负荷和自然条件基础上，采用单母线分段的220kV主接线，"
        "35kV和10kV均采用单母线两分段，0.4kV采用单母线分段暗备用。配置2×180MVA、220/35kV、"
        "YNd11、uk=14%的主变，2×31.5MVA、35/10.5kV、YNd11、uk=8%的T10，2×±12Mvar SVG，"
        "以及2×200kVA、SCB14、Dyn11、uk=4%的所用变。"
        "计算得到35kV汇集计算容量280.286MVA，主变正常负载率77.857%；主变N-1时需限发35.780%。"
        "课程短路模型下10kV条件性最大综合短路电流为15.638kA，禁止并列敏感性为28.472kA。"
        "设备按额定电压、持续电流、开断、峰值和1.10s热稳定进行等级预筛，并补齐CT/PT、零序CT、母线、"
        "绝缘子/套管、MOA、接地变+低电阻和接地开关的课程配置。35/10kV母联转供时仅保留健康侧1套接地源。"
        "CAD按三张必交图加一张L1线路间隔增强详图组织。",
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    set_east_asia_font(p.add_run("关键词："), name="黑体", size=12, bold=True)
    set_east_asia_font(
        p.add_run("新能源汇集变电所；电气主接线；主变压器；短路电流；设备选择；配电装置"),
        name="宋体",
        size=11.5,
    )
    doc.add_page_break()
    add_toc(doc)

    main_sec = start_main_section(doc, SHORT_TITLE + "—技术设计说明书")

    add_heading(doc, "第一章 电气主接线设计", 1, page_break_before=False)
    add_heading(doc, "1.1 原始资料与设计范围", 2)
    add_body(
        doc,
        "设计对象为第5组220kV新能源汇集变电所。本站通过12回35kV集电线路汇集风电、光伏和储能电能，"
        "经两台220/35kV主变升压后由两回220kV线路送出，并预留第三回线路。10kV系统承担站用备用电源、"
        "无功补偿、冷却和集控通信等辅助负荷，0.4kV系统向站内低压负荷供电。",
    )
    add_body(
        doc,
        "本设计的工作链条由原始资料整理、运行方式冻结、负荷与容量计算、短路等值、设备职责确定、"
        "配电装置布置和图纸校核组成。说明书重点回答“为什么采用该方案、方案在什么条件下成立”，"
        "计算书则给出公式、代入和数值校验；两者均读取同一组YAML参数及计算结果，避免表格、图纸与正文各自取值。",
    )
    add_body(
        doc,
        "参数采用分级证据管理：任务书和教师口径优先；现行标准及正式厂家资料用于补充；仍缺少唯一值时，"
        "采用偏保守的课程设计假设并在边界章节中单列。海拔不超过1000m、污秽d级、保护及开断时间、"
        "SVG容量和接地设备参数均属于可被后续真实资料覆盖的课程冻结值。",
    )
    add_table(
        doc,
        "表1-1 主要原始资料",
        ["项目", "任务书或冻结值", "设计处理"],
        [
            ["电压等级", "220/35/10/0.4kV", "覆盖一次主接线、辅助系统和所用电"],
            ["主变压器", "2×180MVA、YNd11、uk=14%", "220/35kV双绕组、油浸、户外、有载调压"],
            ["220kV线路", "2回在运、1回预留；LGJ-400/50", "向西北出线；N-1校核温度修正载流量"],
            ["35kV集电线路", "12回在运、2回预留", "WA/WB共7回架空并设入口MOA；PVA/PVB/ES共5回电缆并设ZCT"],
            ["10kV电源变", "2×31.5MVA、YNd11、uk=8%", "35/10.5kV，分别供10kV-I/II段"],
            ["所用变", "2×200kVA、SCB14、Dyn11、uk=4%", "10/0.4kV，轮换暗备用"],
            ["系统等值", "S1=2400MVA，xc1=0.40；S2=2000MVA，xc2=0.45", "换算至100MVA基准"],
            ["220kV线路长度", "L1=70km，L2=85km，L3=60km", "L3不参加本期基准计算"],
            ["自然条件", "最高41℃，最低-12℃，平均15℃", "设备和导体环境校核"],
            ["场址补充假设", "海拔≤1000m，污秽d级", "课程假设；真实工程须替换"],
        ],
        [3.2, 6.4, 5.9],
        9.3,
    )
    add_result_analysis(
        doc,
        "原始资料决定了本站属于“多回35kV电源汇集、少回220kV集中送出”的典型结构。"
        "主变容量、220kV线路N-1载流量和35/10kV母联事故转供能力因此成为全设计的三条控制主线；"
        "第三回220kV线路仅作远期预留，不参加本期负荷和短路基准计算。",
    )
    add_note_box(
        doc,
        "设计边界",
        "计算、设备等级和CAD几何用于课程设计闭环。精确设备型号、厂家41℃服务条件、完整标准短路模型、"
        "CT/PT/ZCT完整负荷与暂态参数、35/10kV母线机械、MOA能量/TOV、接地源厂家联锁、避雷针保护范围和"
        "接地网接触/跨步电压不在本次课程底稿中冒充已完成。",
    )

    add_heading(doc, "1.2 主接线设计原则", 2)
    principles_paragraph = add_body(
        doc,
        "主接线应满足安全可靠、运行灵活、检修方便、经济合理和便于扩建等要求。本站以新能源集中送出为主，"
        "正常运行时尽量保持各电源分列，事故转供前先隔离故障源，避免通过母联形成未经许可的系统并列。",
    )
    principles_paragraph.paragraph_format.keep_together = True
    add_body(
        doc,
        "可靠性评价不只看断路器数量，还要检查任一线路、主变或母线段退出后的停电范围和恢复路径。"
        "本期方案允许通过母联进行事故转供，但转供能力受存活主变或T10额定容量、线路载流量、短路水平及"
        "接地源联锁共同约束，不能把“母联可以闭合”等同于“所有负荷均可无条件恢复”。",
    )
    add_body(
        doc,
        "经济性评价以满足本期回路规模和远期扩建为前提，避免为有限回路数配置过多母线、隔离开关和占地。"
        "运行灵活性则通过分段、预留间隔、明确的正常断开点和可执行的事故操作顺序实现，而不是依赖长期并列运行。",
    )
    add_bullets(
        doc,
        [
            "220kV侧应适应两回本期线路、两台主变和一回远期预留线路。",
            "35kV侧应把同一场站回路分散到两段母线，并为两台T10和预留回路留出间隔。",
            "10kV侧应保证两套SVG和重要辅助负荷分段供电，禁止两台健康T10长期并列。",
            "0.4kV所用电采用两台所用变暗备用，任一台应能承担全部经常负荷。",
        ],
    )

    add_heading(doc, "1.3 220kV主接线方案比较", 2)
    add_body(
        doc,
        "220kV侧本期包括两回线路、两台主变和一回远期预留线路。候选接线必须同时适应线路与主变成组运行、"
        "母线分区故障、设备检修以及第三回线路扩建。比较时不单纯追求断路器最少，而是综合回路规模、故障影响、"
        "倒闸复杂度、占地投资和课程图纸可表达性。",
    )
    add_table(
        doc,
        "表1-2 220kV候选方案技术比较",
        ["方案", "主要优点", "主要不足", "结论"],
        [
            ["内桥接线", "断路器数量少，初投资低", "远期第三回线路扩建困难；运行灵活性较差", "仅作比较"],
            ["单母线分段", "接线清晰、经济；母线故障限制在一段；预留扩建方便", "一段检修时该段线路和主变退出", "采用"],
            ["双母线不带旁路", "检修灵活、供电可靠、扩建方便", "隔离开关多、占地及投资高、倒闸复杂", "可靠性要求提高时升级"],
        ],
        [3.4, 4.6, 4.6, 2.9],
        9.2,
    )
    scheme_conclusion = add_body(
        doc,
        "综合本期回路数、课程图纸表达和经济性，220kV采用单母线分段。L1、T1接I段，L2、T2接II段，"
        "L3按远期预留间隔设置。分段断路器正常断开；只有调度许可且同期检查通过后才允许条件闭合。",
    )
    scheme_conclusion.paragraph_format.keep_together = True
    add_body(
        doc,
        "内桥接线虽然断路器数量少、初投资较低，但新增第三回线路时桥臂和运行方式均需较大调整，且线路、"
        "主变之间的操作关联较强。双母线方案检修灵活，却会显著增加隔离开关、母线构架、倒闸步骤和占地，"
        "与本期有限回路数不匹配。单母线分段使每段形成“1回线路+1台主变”的清晰单元，母线故障影响限制在一段，"
        "同时为L3保留独立间隔，是本期可靠性与经济性的折中方案。",
    )
    add_result_analysis(
        doc,
        "正常分段运行时系统1和系统2不经本站并列；分段断路器闭合只作为调度许可下的条件方式。"
        "因此设备短路预筛既保留分列运行职责，也对分段闭合场景进行条件性校核，但不把该场景写成默认运行方式。",
    )

    add_heading(doc, "1.4 35kV、10kV与0.4kV接线", 2)
    add_body(
        doc,
        "35kV侧回路数量最多，是本站一次接线组织的重点。两台主变分别接入I、II段母线，12回本期集电线路按6+6分配，"
        "同一风电或光伏场站的多回线路尽量跨段布置，以降低单段母线故障导致整个场站同时退出的概率。"
        "每段另设一回T10高压侧馈线和一回远期预留间隔，母联正常断开。",
    )
    add_body(
        doc,
        "10kV系统由两台35/10.5kV T10分别供电，每段配置一套SVG及相应辅助馈线。采用独立T10而不从主变第三绕组直接取电，"
        "可使10kV无功补偿和辅助系统与主变220/35kV主送出功能分开组织，并便于分段检修；代价是增加两台电源变及相应间隔。"
        "0.4kV系统采用两台所用变轮换暗备用，正常一台运行、母联闭合、另一台进线断开。",
    )
    add_table(
        doc,
        "表1-3 各电压等级接线与正常状态",
        ["电压等级", "接线方式", "正常分段/母联状态", "运行要点"],
        [
            ["220kV", "单母线分段", "断开", "系统1、系统2分列；条件闭合须调度许可和同期"],
            ["35kV", "单母线两分段", "断开", "两台主变各带一段；事故转供前隔离故障主变"],
            ["10kV", "单母线两分段", "断开", "两台T10各带一段；健康T10禁止并列"],
            ["0.4kV", "单母线分段", "闭合", "正常仅一台所用变运行，另一台暗备用"],
        ],
        [2.4, 4.0, 3.1, 6.0],
        9.3,
    )
    add_table(
        doc,
        "表1-4 35kV馈线分配",
        ["母线段", "本期馈线", "有功容量/MW", "附加间隔"],
        [
            ["35kV-I", "WA-1、WA-2、WB-1、WB-2（架空）；PVA-1、PVB-1（电缆）", "130", "T10-1-HV；预留1回"],
            ["35kV-II", "WA-3、WA-4、WB-3（架空）；PVA-2、PVB-2、ES-1（电缆）", "140", "T10-2-HV；预留1回"],
        ],
        [2.6, 7.2, 2.4, 3.3],
        9.2,
    )
    add_body(
        doc,
        "35kV-I段接入130MW，II段接入140MW，两段本期馈线数量相同但负荷构成不同。按基础单回电流累加，"
        "I、II段分别为2230.732A和2404.371A；主变进线及母联设备仍按单台180MVA主变额定电流的1.05倍，"
        "即3117.691A确定持续职责，从而覆盖事故转供而不是只覆盖正常分段负荷。",
    )
    add_result_analysis(
        doc,
        "分段接线的价值在于把故障影响和检修范围局部化。其成立条件是母联不用于两台健康电源长期并列，"
        "并且事故转供前完成故障隔离、容量核对和接地源切换。主接线图中的常开点、预留间隔和接地设备因此都属于运行逻辑的一部分。",
    )
    add_figure(
        doc,
        ROOT / "drawings" / "exports" / "single_line_a1.png",
        "图1-1 SLD-01 220kV新能源汇集变电所电气主接线简图（A1，NTS，课程设计底图）",
    )

    add_heading(doc, "1.5 正常、条件与禁止运行方式", 2, page_break_before=False)
    add_body(
        doc,
        "运行方式是后续持续电流和短路电流选取的前提。正常方式强调电源分列、故障影响受控；条件方式用于事故转供或"
        "系统许可下的临时操作；禁止方式只保留为误操作敏感性和联锁设计依据，不得用来提高日常设备负荷能力。",
    )
    add_table(
        doc,
        "表1-5 运行方式门控",
        ["类别", "运行方式", "处理原则"],
        [
            ["允许", "220/35/10kV分段分列；0.4kV单台所用变带两段", "正常运行基准"],
            ["条件允许", "220kV分段闭合", "调度许可并同期检查后方可闭合"],
            ["事故转供", "一台主变或一台T10退出后闭合下一级母联", "先隔离故障源；退出受电段接地源，仅保留健康侧1套"],
            ["禁止", "两台健康主变经35kV母联并列", "仅保留短路敏感性，不作为运行方式"],
            ["禁止", "两台健康T10经10kV母联并列", "联锁闭锁；避免短路电流显著上升"],
            ["禁止", "35/10kV母联闭合且两套接地源同时投入", "低电阻并联会改变接地故障电流与保护配合"],
            ["禁止", "两台所用变经0.4kV并列", "两进线与母联设置电气、机械联锁"],
        ],
        [2.8, 5.8, 6.9],
        9.2,
    )
    add_body(doc, "发生主变或T10故障时，事故转供按下列顺序进行：")
    add_bullets(
        doc,
        [
            "保护动作切除故障设备，并确认相应断路器、隔离开关已形成可靠电气隔离。",
            "核对健康电源、母线和联络回路的电压、容量、持续电流及短路职责，确认转供后不超过本报告的控制值。",
            "35kV或10kV母联合闸前，先退出受电侧或故障侧接地变+低电阻，仅保留健康电源侧1套接地源。",
            "满足同期、闭锁和调度许可后闭合母联；主变N-1时同步执行新能源限发，使存活主变负荷不超过180MVA。",
            "故障设备恢复、母联重新断开后，再恢复每个独立带电母线段各1套接地源。",
        ],
        ordered=True,
    )
    add_result_analysis(
        doc,
        "两套低电阻接地源经母联并联会降低等效接地电阻、改变目标单相接地电流和保护配合，"
        "因此“母联已闭合且两套接地源同时投入”被明确列为禁止状态。相同地，两台健康主变或T10并列会显著抬高短路电流，"
        "两台所用变并列还可能形成低压环流，均应由电气和机械联锁阻止。",
    )

    add_heading(doc, "第二章 负荷计算及变压器选择", 1)
    add_heading(doc, "2.1 35kV与10kV综合负荷", 2)
    add_body(
        doc,
        "35kV新能源负荷按各项视在功率求和，再计0.95最大同时出力系数和5%线损；单回馈线电流按各类"
        "最大负荷在回路间平均分配，不重复乘同时系数。10kV辅助负荷按0.80同时系数和5%线损计算。",
    )
    add_equation(doc, "Sᵢ = Pᵢ / cosφᵢ", "2-1")
    add_equation(doc, "I = S / (√3 U)", "2-2")
    add_formula_terms(
        doc,
        "Pᵢ、Sᵢ分别为第i类负荷的有功功率和视在功率，cosφᵢ为功率因数；U为三相线电压，"
        "当S以MVA、U以kV代入时，I的结果为kA，换算为A需乘1000。",
    )
    add_body(
        doc,
        "35kV五类负荷的视在功率分别为84.210526、63.157895、61.224490、40.816327和31.578947MVA，"
        f"标量合计为{load35['gross_apparent_mva']:.6f}MVA。主变容量按下式计入同时出力和线损：",
    )
    add_equation(
        doc,
        f"S35 = {load35['gross_apparent_mva']:.6f} × {load35['group_factor']:.2f} × "
        f"(1 + {load35['loss_rate']:.2f}) = {load35['with_losses_mva']:.6f} MVA",
        "2-3",
    )
    add_body(
        doc,
        f"10kV三类辅助负荷原始视在功率合计{load10['gross_apparent_mva']:.6f}MVA，"
        f"按0.80同时系数和5%线损得到{load10['with_losses_mva']:.6f}MVA。"
        "该负荷在主变主校核中不与35kV汇集容量重复叠加，只作为保守敏感性保留。",
    )
    add_table(
        doc,
        "表2-1 综合负荷计算结果",
        ["项目", "结果", "说明"],
        [
            ["35kV原始视在功率", "280.988MVA", "270MW按逐项功率因数折算"],
            ["计0.95同时出力", "266.939MVA", "主变容量计算用"],
            ["再计5%线损", "280.286MVA", "本期主变校核容量"],
            ["10kV辅助负荷", "1.757MVA", "计0.80同时系数和5%线损"],
            ["保守叠加10kV敏感性", "282.043MVA", "不替代主校核口径"],
        ],
        [5.0, 3.6, 6.9],
        9.5,
    )
    add_result_analysis(
        doc,
        f"将各项视在功率标量相加得到{load35['with_losses_mva']:.6f}MVA，属于偏保守的课程口径；"
        f"若先分别合成P、Q再求模，结果为{load35['vector_with_losses_mva']:.6f}MVA，"
        f"两者相差{load35['with_losses_mva'] - load35['vector_with_losses_mva']:.6f}MVA。"
        "为与任务书和既有容量计算保持一致，主变选择采用前者，向量结果只用于数量级复核。",
    )

    add_heading(doc, "2.2 主变压器选择", 2)
    add_body(
        doc,
        "选用2×180MVA、220/35kV双绕组油浸式有载调压主变，联结组YNd11，短路电压uk=14%。"
        "220kV星形中性点经专用中性点CT直接接地，35kV侧为三角形绕组。两台正常运行时"
        "每台承担140.143MVA，负载率77.857%。一台退出时，另一台额定容量只能覆盖64.220%的最大计算负荷，"
        "缺口100.286MVA，因此按任务书允许的新能源限发策略，基准限发比例为35.780%。",
    )
    add_body(
        doc,
        "主变采用220/35kV双绕组方案，并以两台独立35/10.5kV T10承担10kV电源功能。与三绕组主变相比，"
        "该方案使10kV辅助系统和SVG运行方式更独立，便于两段母线分别检修与控制；同时会增加T10及其35kV馈线间隔。"
        "结合任务书已给的2×180MVA容量和当前辅助系统规模，本阶段采用双绕组主变加T10方案。",
    )
    add_equation(
        doc,
        f"Snormal,each = {main_transformer['total_required_mva']:.6f} / 2 = "
        f"{main_transformer['per_transformer_mva']:.6f} MVA",
        "2-4",
    )
    add_equation(
        doc,
        f"βnormal = {main_transformer['per_transformer_mva']:.6f} / 180 × 100% = "
        f"{main_transformer['normal_loading_percent']:.3f}%",
        "2-5",
    )
    add_equation(
        doc,
        f"KN-1 = 180 / {main_transformer['total_required_mva']:.6f} × 100% = "
        f"{main_transformer['n_minus_one_supply_percent']:.3f}%",
        "2-6",
    )
    add_equation(
        doc,
        f"Kcurtail = (1 − 180 / {main_transformer['total_required_mva']:.6f}) × 100% = "
        f"{100 - main_transformer['n_minus_one_supply_percent']:.3f}%",
        "2-7",
    )
    add_table(
        doc,
        "表2-2 主变容量与N-1校验",
        ["指标", "计算值", "结论"],
        [
            ["装机容量", "360.000MVA", "2×180MVA"],
            ["正常每台负荷", "140.143MVA", "均衡分配"],
            ["正常负载率", "77.857%", "满足课程设计经济运行要求"],
            ["N-1覆盖比例", "64.220%", "低于全部最大负荷"],
            ["N-1容量缺口", "100.286MVA", "通过新能源限发处理"],
            ["N-1限发比例", "35.780%", "当前总N-1控制因素"],
        ],
        [5.0, 4.0, 6.5],
        9.5,
    )
    add_result_analysis(
        doc,
        "2×180MVA能够满足正常方式，但不能在不采取控制措施时满足全部最大出力下的主变N-1。"
        "方案成立的运行条件是新能源场站具备快速限发能力，并在事故转供前把存活主变负荷限制在180MVA以内。"
        "1.05系数仅用于主变回路设备持续电流预筛，不表示主变本体可以长期按189MVA运行。",
    )

    add_heading(doc, "2.3 10kV电源变与无功补偿", 2)
    add_body(
        doc,
        "10kV由2×31.5MVA、35/10.5kV、YNd11、uk=8%的T10供电，每段配置±12Mvar SVG。"
        "按目标功率因数0.98，需补偿21.8875Mvar；选24Mvar后裕度9.65%，pf≈0.98147。"
        "T10正常负载率39.58%，N-1最严吸收工况79.17%。",
    )
    add_body(
        doc,
        "无功补偿按不抵扣新能源变流器自身无功裕度的保守口径计算。35kV汇集负荷与10kV辅助负荷计损耗后，"
        f"合计有功P={reactive['basis']['p35_mw'] + reactive['basis']['p10_mw']:.3f}MW，"
        f"无功Q={reactive['basis']['q35_mvar'] + reactive['basis']['q10_mvar']:.6f}Mvar。"
        "目标功率因数对应的允许无功为P·tan(arccos0.98)，两者之差即所需补偿容量。",
    )
    add_equation(doc, "Qc = Q − P tan(arccos cosφtarget)", "2-8")
    add_equation(
        doc,
        f"Qc = {reactive['basis']['q35_mvar'] + reactive['basis']['q10_mvar']:.6f} − "
        f"{reactive['basis']['p35_mw'] + reactive['basis']['p10_mw']:.3f} × tan(arccos0.98) "
        f"= {reactive['calculated_required_mvar_conservative_with_10kv_auxiliary']:.6f} Mvar",
        "2-9",
    )
    add_equation(
        doc,
        f"Kmargin = (24 / {reactive['calculated_required_mvar_conservative_with_10kv_auxiliary']:.6f} − 1) × 100% "
        f"= {reactive['selected_margin_percent']:.2f}%",
        "2-10",
    )
    add_body(
        doc,
        f"每台T10正常均衡承担{reactive['source_transformers']['normal_balanced_mva_each']:.6f}MVA，"
        f"负载率{reactive['source_transformers']['normal_balanced_loading_percent']:.3f}%；"
        f"N-1全吸收无功工况为{reactive['source_transformers']['n_minus_one_full_absorbing_mva']:.6f}MVA，"
        f"负载率{reactive['source_transformers']['n_minus_one_full_absorbing_loading_percent']:.3f}%。"
        "因此2×31.5MVA既覆盖正常分段运行，也为单台T10事故承担两套SVG的最严课程场景保留约20%的容量余量。",
    )
    add_table(
        doc,
        "表2-3 T10与SVG配置",
        ["设备", "配置", "正常状态", "校验结论"],
        [
            ["T10-1、T10-2", "2×31.5MVA，35/10.5kV，YNd11，uk=8%", "各带一段10kV母线", "N-1最严79.17%"],
            ["SVG-1、SVG-2", "2×±12Mvar", "每段一套，电压/无功/功率因数控制", "补偿后pf≈0.98147"],
        ],
        [3.7, 4.9, 3.8, 3.1],
        9.2,
        "容量属于任务书缺项下的课程冻结值，最终由无功电压与谐波专题覆盖。",
    )
    add_result_analysis(
        doc,
        "两套SVG分置于10kV-I、II段可在正常方式下就地调节电压、无功或功率因数。"
        "24Mvar只比计算需求高9.65%，属于合理但不宽裕的课程选择；最终容量仍须由并网点电压偏差、"
        "动态响应、谐波和变流器持续过载能力专题复核。",
    )

    add_heading(doc, "2.4 所用变压器选择", 2)
    station_service_paragraph = add_body(
        doc,
        "扣除110kV不适用项后，连续负荷96.3kW，最严场景需192.913kVA。"
        "选用2×200kVA、10/0.4kV、SCB14干式、Dyn11、uk=4%的所用变，每台可带全部经常负荷并留约10%裕度。"
        "正常一台供电且0.4kV母联闭合，另一台及进线暗备用。",
    )
    station_service_paragraph.paragraph_format.keep_together = True
    add_body(
        doc,
        "所用电负荷按连续、经常短时和不经常短时分类。基准场景取全部连续负荷96.3kW加空气压缩机同时投入值22.0kW，"
        "合计118.3kW；最严经常场景把两台空气压缩机44.0kW全部计入，合计140.3kW。"
        "110kV配电装置和110kV断路器加热负荷因本项目无该电压等级而明确剔除。",
    )
    add_equation(doc, "ST = Pscenario (1 + Km) / cosφ", "2-11")
    add_formula_terms(doc, "Km为容量裕度，本设计取10%；所用电综合功率因数按课程假设取0.80。")
    add_equation(
        doc,
        f"Sbase = {station_service['base_scenario_kw_before_margin']:.1f} × 1.10 / 0.80 = "
        f"{station_service['base_required_kva']:.4f} kVA",
        "2-12",
    )
    add_equation(
        doc,
        f"Sworst = {station_service['worst_frequent_kw_before_margin']:.1f} × 1.10 / 0.80 = "
        f"{station_service['worst_frequent_required_kva']:.4f} kVA",
        "2-13",
    )
    add_result_analysis(
        doc,
        "按最严场景选择200kVA标准容量，每台所用变均可单独承担全部课程负荷。暗备用方式避免两台变压器并列，"
        "同时通过母联闭合保证正常仅一台运行即可带两段低压母线；备用投入前仍须先切除故障进线并满足防并列联锁。",
    )

    add_heading(doc, "第三章 最大持续工作电流及短路计算", 1)
    add_heading(doc, "3.1 各回路最大持续工作电流", 2)
    add_body(
        doc,
        "最大持续工作电流用于确定断路器、隔离开关、CT、母线和开关柜的额定电流。馈线按本类最大出力、实际回路数和5%线损计算，"
        "不再乘全站0.95同时出力系数；主变、T10和SVG回路则按设备额定容量乘1.05选型裕度确定。不同口径服务于不同对象，不能相互替代。",
    )
    add_equation(doc, "Ifeeder = P (1 + kloss) / (n √3 U cosφ)", "3-1")
    add_formula_terms(doc, "n为同类负荷回路数，kloss=0.05；馈线计算反映单回线路可能承担的本类最大出力。")
    add_equation(
        doc,
        "IES = 30 × 1.05 / (1 × √3 × 35 × 0.95) × 1000 = 546.963 A",
        "3-2",
    )
    add_body(
        doc,
        "储能馈线只有1回，因而形成35kV单回馈线最大电流546.963A；风电A、B虽总容量不同，但80MW/4回与60MW/3回的单回有功均为20MW，"
        "故两者单回设计电流相同，均为364.642A。",
    )
    feeder_rows = []
    for item in load["load_35kv"]["items"]:
        feeder_rows.append(
            [
                item["label_zh"],
                item["circuits"],
                f'{item["per_circuit_current_a"]:.3f}',
                "计5%线损，不乘0.95",
            ]
        )
    add_table(
        doc,
        "表3-1 35kV单回馈线设计电流",
        ["负荷类别", "回路数", "单回电流/A", "计算口径"],
        feeder_rows,
        [4.2, 2.2, 3.4, 5.7],
        9.5,
    )
    add_table(
        doc,
        "表3-2 主要回路持续电流职责",
        ["回路", "最大持续电流/A", "设计说明"],
        [
            ["220kV线路间隔", "735.559", "一回退出时另一回承担全部本期送出"],
            ["220kV分段", "735.559", "一段事故转供全部本期送出功率的课程保守职责"],
            ["主变220kV侧", "495.996", "180MVA额定电流×1.05"],
            ["主变35kV侧/35kV母联", "3117.691", "180MVA额定电流×1.05"],
            ["35kV-I段", "2230.732", "按I段实际6回馈线累加"],
            ["35kV-II段", "2404.371", "按II段实际6回馈线累加"],
            ["T10 35kV侧", "545.596", "31.5MVA额定电流×1.05"],
            ["T10 10kV设备口径", "1909.586", "按10kV开关设备电压基准×1.05"],
            ["单套SVG馈线", "692.820", "12Mvar、10.5kV额定电流×1.05"],
        ],
        [5.0, 3.7, 6.8],
        9.2,
    )
    add_body(
        doc,
        "35kV母线段存在三个容易混淆的电流口径。按未计线损的各馈线基础电流和实际分配累加，I、II段分别为2230.732A和2404.371A；"
        "若按各馈线基础视在功率标量求和后统一施加0.95同时系数和5%线损，则分别约为2225.155A和2398.360A；"
        "按各段P、Q向量合成则分别为2221.564A和2394.612A。若把每回馈线均按自身最大值计5%线损后直接相加，"
        "则分别约为2342.268A和2524.590A。前三者用于不同的正常负荷分析，最后一种是各馈线同时满出的保守和。",
    )
    add_result_analysis(
        doc,
        "主变进线、35kV母联及35kV主母线的最终设备持续职责不采用上述较小的正常段电流，而按单台主变额定电流1.05倍取3117.691A，"
        "以覆盖事故转供。35kV 3150A开关柜相对该职责仅余32.309A，后续必须维持室内温度不高于40℃并取得厂家温升确认。",
    )

    add_heading(doc, "3.2 220kV线路导线校核与N-1", 2)
    add_body(
        doc,
        "任务书指定220kV送出导线为LGJ-400/50。课程参考载流量592A对应环境25℃、导体允许温度70℃，"
        "本站最高环境温度为41℃，因此必须先进行温度修正，再分别与正常双回和单回事故电流比较。",
    )
    add_equation(doc, "KT = √[(θallow − θproject)/(θallow − θreference)]", "3-3")
    add_equation(doc, "KT = √[(70 − 41)/(70 − 25)] = 0.802773", "3-4")
    add_equation(doc, "Iallow,41 = 592 × 0.802773 = 475.242 A", "3-5")
    conductor_paragraph = add_body(
        doc,
        "LGJ-400/50课程参考载流量为592A（25℃环境、70℃导体）。41℃修正系数0.802773，课程允许值"
        "475.242A；正常每回367.780A通过。单回全送735.559A时需限发35.390%；主变N-1限发后线路电流"
        "472.377A，余量2.864A，因此当前仍由主变容量控制。",
    )
    conductor_paragraph.paragraph_format.keep_together = True
    add_equation(
        doc,
        f"Kline-curtail = (1 − {outgoing_ampacity['corrected_ampacity_a']:.3f} / "
        f"{outgoing['single_circuit_contingency_current_a']:.3f}) × 100% = "
        f"{outgoing_ampacity['line_only_required_curtailment_percent']:.3f}%",
        "3-6",
    )
    add_result_analysis(
        doc,
        f"主变N-1限发后，220kV线路电流为{outgoing_ampacity['main_transformer_n_minus_one_current_a']:.3f}A，"
        f"比课程允许值只低{outgoing_ampacity['margin_after_main_transformer_n_minus_one_a']:.3f}A。"
        "主变限发比例虽然略严于导线单独约束，但线路热裕度非常小；真实工程必须用明确风速、日照、环境和导体温度条件下的厂家热额定值覆盖本课程参考。",
    )
    add_note_box(
        doc,
        "证据边界",
        "592A来自旧版二手课程参数表，仅作保守校核；厂家热额定值须结合风速、日照、环境及导体温度复核。",
    )

    add_heading(doc, "3.3 短路计算方法与计算点", 2)
    add_body(
        doc,
        "短路初算采用任务书规定的100MVA基准容量和230/37/10.5kV平均额定电压。网络按标幺电抗法处理，"
        "初步忽略电阻；主变电抗按uk=14%，T10按uk=8%。新能源和SVG贡献在敏感性中按额定电流1.1~1.2倍"
        "算术叠加，峰值采用课程系数k=1.8。该方法用于课程设备等级预筛，不替代GB/T 15544.1-2023完整计算。",
    )
    add_body(
        doc,
        "计算时先把系统、线路和变压器电抗统一换算到100MVA基准，再按实际运行方式进行串并联组合。"
        "220kV分列时每段只有对应系统和线路路径；220kV分段条件闭合时两条路径并联；向35kV、10kV计算点继续串入主变和T10电抗。",
    )
    add_equation(doc, "X*sys = xc × SB / Ssys", "3-7", compact=True)
    add_equation(doc, "X*line = xL × SB / UB²", "3-8", compact=True)
    add_equation(doc, "X*T = (uk% / 100) × SB / SN", "3-9", compact=True)
    add_equation(doc, "Ik = SB / (√3 UB XΣ*)", "3-10", compact=True)
    add_equation(doc, "ip = √2 k Ik", "3-11", compact=True)
    add_formula_terms(
        doc,
        "xc为系统在自身容量基准上的标幺电抗，x为线路单位长度电抗，L为线路长度，uk为变压器短路电压百分数；"
        "XΣ*为故障点至各电源的等值标幺电抗。",
    )
    add_body(
        doc,
        "系统1、系统2连同L1、L2后的路径电抗分别为0.069597和0.086772，分段闭合后的并联等值为0.038621。"
        f"220kV分段闭合电网分量为{sc_points['SC-220-BUS-CLOSED']['grid_symmetrical_current_ka']:.3f}kA，"
        f"叠加新能源上界后为{sc_points['SC-220-BUS-CLOSED']['conservative_total_symmetrical_current_range_ka']['maximum']:.3f}kA。",
    )
    add_body(
        doc,
        f"35kV条件性最大综合短路电流为{sc_points['SC-35-II-220-CLOSED']['conservative_total_symmetrical_current_range_ka']['maximum']:.3f}kA，"
        f"10kV条件性最大值为{sc_points['SC-10-I-220-CLOSED']['conservative_total_symmetrical_current_range_ka']['maximum']:.3f}kA。"
        f"两台健康主变或两台健康T10并列时分别升至{sc_points['SC-35-BOTH-TRANSFORMERS-SENSITIVITY']['conservative_total_symmetrical_current_range_ka']['maximum']:.3f}kA和"
        f"{sc_points['SC-10-BOTH-T10-SENSITIVITY']['conservative_total_symmetrical_current_range_ka']['maximum']:.3f}kA，只用于联锁和误操作风险提示。",
    )
    add_table(
        doc,
        "表3-3 短路职责分级（禁止并列值仅作联锁与误操作风险提示）",
        ["电压级", "正常必选/kA", "条件预校核/kA", "禁并敏感/kA", "课程峰值/kA"],
        [
            ["220kV", "4.033", "7.385", "—", "18.798"],
            ["35kV", "13.265", "16.291", "25.694", "41.470"],
            ["10kV", "14.492", "15.638", "28.472", "39.808"],
        ],
        [2.2, 3.3, 3.5, 3.4, 3.1],
        8.4,
        margin_x=75,
        margin_y=50,
    )
    add_result_analysis(
        doc,
        "开关设备的条件性预校核职责取220/35/10kV的7.385、16.291和15.638kA，"
        "而不是只取正常分列的较小值。禁止并列敏感性不作为日常必选开断职责，但若联锁不能可靠阻止该方式，"
        "就必须重新提高设备等级或改变系统运行方案。",
    )

    add_heading(doc, "第四章 主要电气设备选择", 1)
    add_heading(doc, "4.1 选择与校验原则", 2)
    add_body(
        doc,
        "一次设备按额定电压、最大持续工作电流、短路开断、关合/峰值耐受和热稳定逐项校验。热稳定等值时间"
        "按后备保护1.00s与断路器总开断0.08s向上取1.10s。环境最高温度41℃、海拔≤1000m和污秽d级作为"
        "课程设计条件。室外主变、T10、可能的户外SVG及导体按41℃复核，35/10kV室内开关柜统一按受控≤40℃；"
        "厂家精确服务条件未取得时，设备状态只能写为额定值等级预筛通过。",
    )
    add_equation(doc, "Ue ≥ Un；Ie ≥ Imax", "4-1")
    add_equation(doc, "Ibr ≥ Ik；ipeak,rated ≥ ip", "4-2")
    add_equation(doc, "Ith² tr ≥ Ik² t", "4-3")
    add_formula_terms(
        doc,
        "Ue、Ie为设备额定电压和额定电流；Ibr为断路器额定开断电流；ipeak,rated为额定峰值耐受或关合能力；"
        "Ith、tr为设备短时耐受电流及其额定持续时间。",
    )
    add_body(
        doc,
        "设备选择按“先确定回路职责、再选择额定值等级、最后逐项比较”的顺序进行。持续电流取正常或事故方式中的较大值；"
        "短路开断和峰值采用条件性最大课程职责；热稳定等值时间由1.00s后备保护和0.08s全开断时间相加后向上取1.10s。"
        "同一设备只有所有适用数值条件满足，且未完成的厂家/环境项目被明确保留时，才能写为“课程预筛通过”。",
    )

    add_heading(doc, "4.2 断路器、隔离开关和开关柜", 2)
    add_table(
        doc,
        "表4-1 主要开关设备课程目标等级",
        ["安装位置", "目标等级", "额定电流/A", "开断/kA", "峰值/kA", "热稳定"],
        [
            ["220kV线路/主变/分段断路器", "252kV户外SF6", "3150", "50", "125", "50kA/3s"],
            ["220kV隔离开关", "252kV户外", "3150", "—", "125", "50kA/3s"],
            ["35kV主变进线/母联柜", "40.5kV金属铠装", "3150", "31.5", "80", "31.5kA/4s"],
            ["35kV馈线/T10柜", "40.5kV金属铠装", "1250", "31.5", "80", "31.5kA/4s"],
            ["35kV接地变馈线柜", "40.5kV金属铠装", "1250", "31.5", "80", "31.5kA/4s"],
            ["10kV进线/母联柜", "12kV金属铠装", "2500", "31.5", "80", "31.5kA/4s"],
            ["10kV SVG/辅助馈线柜", "12kV金属铠装", "1250", "31.5", "80", "31.5kA/4s"],
            ["10kV接地变馈线柜", "12kV金属铠装", "1250", "31.5", "80", "31.5kA/4s"],
        ],
        [4.0, 2.9, 2.0, 1.8, 1.8, 3.0],
        8.8,
        "35kV 3150A相对3117.691A仅余32.309A（约1.04%）；10kV已升级为2500A，相对1909.586A"
        "余590.414A。接地变馈线分别覆盖400A/200A课程职责。35/10kV柜统一按室内受控≤40℃，仍须由厂家温升和精确柜型确认。",
        nowrap_header_columns={2, 3, 4, 5},
        nowrap_body_columns={2, 3, 4, 5},
    )
    add_body(
        doc,
        "以35kV主变进线柜为例，额定电流3150A与职责3117.691A相比仅有32.309A裕度，约为1.04%；"
        "额定开断电流31.5kA高于条件性短路职责16.291kA，额定峰值耐受80kA高于41.470kA。"
        "热稳定比较为31.5²×4=3969kA²·s，大于16.291²×1.10=291.936kA²·s。",
    )
    add_body(
        doc,
        "10kV进线和母联采用2500A而非较低等级，是因为T10按10kV设备口径并计1.05裕度后的职责达到1909.586A。"
        "2500A方案保留590.414A裕度，31.5kA开断和80kA峰值等级也分别覆盖15.638kA和39.808kA课程职责。",
    )
    add_result_analysis(
        doc,
        "35kV 3150A等级的短路和热稳定裕度较大，但持续电流裕度很小，故其通过结论必须附带“开关室≤40℃、"
        "通风空调可靠、厂家温升试验和精确柜型确认”的条件。不能因短路裕度充足而忽略长期发热控制。",
    )

    add_table(
        doc,
        "表4-2 已审查设备的适用性结论",
        ["设备", "课程结论", "依据与复核条件"],
        [
            ["限流电抗器", "35/10kV均不设置", "16.291/15.638kA低于31.5kA设备能力；最终短路职责超过31.5kA时重审"],
            ["高压熔断器", "不作为主回路设备", "仅在35/10kV PT柜按厂家方案设置一次熔断器及二次保护"],
            ["阻波器", "当前不配置", "仅在通信专业确定电力线载波通信及耦合方案时配置"],
        ],
        [3.0, 4.0, 8.5],
        9.0,
    )
    add_body(
        doc,
        "负面设备选择同样需要给出计算依据。35kV和10kV条件性短路职责分别为16.291kA和15.638kA，"
        "均明显低于所选31.5kA开关设备能力，因此本课程方案不设置限流电抗器。高压熔断器只可能作为PT柜成套保护元件，"
        "不承担主回路开断；阻波器则取决于通信专业是否采用电力线载波，当前任务书未给该条件。",
    )

    add_heading(doc, "4.3 CT与PT/CVT课程配置", 2)
    add_body(
        doc,
        "各电压等级每段主母线配置电压互感器；凡装有断路器的回路原则上配置三相CT，电缆单相接地保护另设"
        "零序CT。下表冻结一次变比、课程芯级组合和职责，但负担、饱和、拐点电压、暂态性能、窗口尺寸、"
        "电缆压降和精确型号仍等待保护计量清册。",
    )
    add_table(
        doc,
        "表4-3 CT课程目标配置",
        ["位置", "变比", "芯级目标", "短时/峰值"],
        [
            ["220kV线路", "1000/1A", "0.2S、0.5、5P30，按需PX", "50kA/3s；125kA"],
            ["220kV主变", "600/1A", "0.2S、0.5、5P30、PX", "50kA/3s；125kA"],
            ["主变220kV中性点", "600/1A", "PX限制接地故障、5P30中性点过流", "单相接地/零序专题待定"],
            ["220kV分段", "1000/1A", "0.5、5P30母线保护、PX母差", "50kA/3s；125kA"],
            ["35kV进线/母联", "4000/1A", "0.2S、0.5、5P30、PX", "31.5kA/4s；80kA"],
            ["35kV馈线/T10", "800/1A", "0.2S、0.5、5P30", "31.5kA/4s；80kA"],
            ["10kV进线/母联", "2500/1A", "0.2S、0.5、5P30、PX", "31.5kA/4s；80kA"],
            ["10kV馈线/SVG", "1000/1A", "0.2S、0.5、5P30", "31.5kA/4s；80kA"],
            ["35kV电缆馈线ZCT", "100/1A", "5P20或厂家等效零序保护级", "5回；一次剩余电流目标400A"],
            ["10kV电缆/SVG ZCT", "50/1A", "5P20或厂家等效零序保护级", "2回SVG+3类辅助；目标200A"],
        ],
        [3.5, 2.6, 5.7, 3.7],
        8.2,
        "短时栏峰值单位为kA；ZCT职责与三相CT分开，三相电缆芯线同穿窗口而屏蔽接地回流线不得穿入。课程芯级组合不作为订货清册。",
    )
    add_body(
        doc,
        "CT一次变比先按最大持续电流选择，再检查短时、峰值及保护用途。35kV进线4000/1A覆盖3117.691A，"
        "10kV进线2500/1A覆盖1909.586A，35kV馈线800/1A覆盖最大单回546.963A；相同电压级内不机械统一变比，"
        "以避免馈线测量分辨率过低。0.2S/0.5级分别面向计量和测量，5P30及PX课程目标用于过流、差动或母差保护。",
    )
    add_body(
        doc,
        "主变220kV中性点CT的持续和短路职责不能沿用主变相回路的三相短路电流。其600/1A、PX+5P30仅为课程配置目标，"
        "最终必须依据单相接地、零序网络、保护负担和拐点电压专题确定。35/10kV电缆回路另设ZCT，"
        "用于剩余电流和单相接地保护，窗口尺寸及屏蔽层接地回流路径需与电缆敷设共同复核。",
    )
    add_table(
        doc,
        "表4-4 PT/CVT课程目标配置",
        ["电压级", "型式与一次电压", "二次绕组", "准确级"],
        [
            ["220kV", "每段单相CVT组；220/√3kV", "100/√3V测量、100/√3V保护、100/3V开口三角", "0.2、3P"],
            ["35kV", "每段三台单相电磁式VT；35/√3kV", "100/√3V测量、100/√3V保护、100/3V开口三角", "0.5、3P"],
            ["10kV", "每段三台单相电磁式VT；10/√3kV", "100/√3V测量、100/√3V保护、100/3V开口三角", "0.5、3P"],
        ],
        [2.2, 4.2, 6.7, 2.4],
        8.5,
        "容量、铁磁谐振、同期抽取和厂家精确型号仍待二次专业复核。",
    )
    add_result_analysis(
        doc,
        "PT/CVT配置满足测量、保护和开口三角零序电压取得的课程功能，但一次/二次电压和准确级并不能代替容量校验。"
        "最终二次负荷、导线压降、铁磁谐振抑制、CVT暂态特性及同期抽取回路应由保护计量清册和厂家资料闭合。",
    )

    add_heading(doc, "4.4 MOA、中性点接地与接地开关", 2)
    add_table(
        doc,
        "表4-5 MOA参数与绝缘裕度课程预筛",
        ["电压级/型号", "Ur/Uc/In/残压", "绝缘裕度", "配置位置"],
        [
            ["220kV YH10W-204/532", "204kV/159kV/10kA/≤532kV", "Uc余13.508kV；950/532=1.786＞1.15", "线路入口、主变高压端"],
            ["35kV YH5WZ-51/134", "51kV/40.8kV/5kA/≤134kV", "Uc仅余0.300kV（约0.74%）；185/134=1.381＞1.15", "母线PT柜、主变端、7回架空入口"],
            ["10kV YH5WZ-17/45", "17kV/13.6kV/5kA/≤45kV", "Uc余1.600kV；75/45=1.667＞1.15", "母线PT柜、T10/SVG端"],
        ],
        [3.5, 4.4, 4.6, 3.0],
        8.2,
        "仅为课程预筛；35kV的TOV、能量、引线压降、线路放电等级和精确型号待专题复核。",
        keep_note_with_table=True,
    )
    add_body(
        doc,
        "MOA课程预筛包括持续运行电压和雷电保护裕度两项。持续运行电压Uc应覆盖相应接地方式下的长期相对地电压，"
        "保护裕度以被保护设备雷电冲击耐受电压LIWV与避雷器残压之比进行初步判断。",
    )
    add_equation(doc, "Uc ≥ Uc,required；Kp = LIWV / Ures", "4-4")
    add_body(
        doc,
        f"220kV设备最高电压252kV，有效接地课程口径下Uc,required=252/√3="
        f"{arresters['MOA-220']['calculated']['required_continuous_voltage_kv']:.3f}kV；"
        f"选用159kV Uc后裕度{arresters['MOA-220']['calculated']['continuous_voltage_margin_kv']:.3f}kV，"
        f"LIWV/残压={arresters['MOA-220']['calculated']['protection_ratio']:.3f}。",
    )
    add_body(
        doc,
        f"35kV系统按非有效接地课程上界取Uc,required="
        f"{arresters['MOA-35']['calculated']['required_continuous_voltage_kv']:.3f}kV，"
        f"所选40.8kV Uc仅余{arresters['MOA-35']['calculated']['continuous_voltage_margin_kv']:.3f}kV，"
        f"约{arresters['MOA-35']['calculated']['continuous_voltage_margin_kv'] / arresters['MOA-35']['calculated']['required_continuous_voltage_kv'] * 100:.2f}%。"
        "该裕度是本组MOA预筛中最薄弱的一项，必须由10s TOV、能量和厂家曲线复核。",
    )
    add_table(
        doc,
        "表4-6 接地参数与接地开关课程配置",
        ["位置", "课程配置", "等级/参数", "工程边界"],
        [
            ["主变220kV中性点", "经600/1A中性点CT直接接地", "PX+5P30课程目标", "持续/短时/峰值按单相接地与零序专题，不沿用三相回路职责"],
            ["35kV-I/II段", "每段1000kVA ZN接地变+低电阻", "400A/10s；R≈50.5Ω", "2套；1250A馈线柜；600/1A相CT+400/1A中性点CT"],
            ["10kV-I/II段", "每段200kVA ZN接地变+低电阻", "200A/10s；R≈28.9Ω", "2套；1250A馈线柜；300/1A相CT+200/1A中性点CT"],
            ["35/10kV母联转供", "合母联前退出受电/故障段接地源", "仅保留健康侧1套", "联锁禁止母联合闸且两套接地源并联；分列后恢复每段1套"],
            ["220kV线路侧ES", "线路侧QS配常开ES", "252kV、50kA/3s、125kA峰值", "感应电流开合等级待厂家专题"],
            ["35/10kV柜内ES", "柜内常开并设机械/电气联锁", "31.5kA/4s、80kA峰值", "精确柜型和联锁图待核"],
        ],
        [3.1, 4.5, 4.2, 3.7],
        8.1,
    )
    grounding_method_paragraph = add_body(
        doc,
        "35kV和10kV绕组均不能直接提供中性点，因此每段采用ZN接地变形成中性点，并串入低电阻限制单相接地电流。"
        "电阻和接地变容量按相电压、目标接地电流及10s短时热容量计算。",
    )
    grounding_method_paragraph.paragraph_format.keep_together = True
    add_equation(doc, "Uφ = Un/√3；R = Uφ/Ig；Seq = Uφ Ig", "4-5")
    add_equation(doc, "ST,min = Seq / Koverload", "4-6")
    add_body(
        doc,
        f"35kV系统Uφ={grounding_packages['GRD-PKG-35']['calculated']['phase_to_earth_voltage_kv']:.4f}kV，"
        f"R={grounding_packages['GRD-PKG-35']['calculated']['resistance_ohm']:.3f}Ω，取50.5Ω；"
        f"10s等效功率{grounding_packages['GRD-PKG-35']['calculated']['short_time_equivalent_power_mva']:.3f}MVA，"
        f"按10倍短时过载折算最小容量{grounding_packages['GRD-PKG-35']['calculated']['minimum_transformer_capacity_kva_at_overload_factor']:.3f}kVA，故选1000kVA。",
    )
    add_body(
        doc,
        f"10kV系统相应得到R={grounding_packages['GRD-PKG-10']['calculated']['resistance_ohm']:.3f}Ω，取28.9Ω；"
        f"10s等效功率{grounding_packages['GRD-PKG-10']['calculated']['short_time_equivalent_power_mva']:.3f}MVA，"
        f"最小接地变容量{grounding_packages['GRD-PKG-10']['calculated']['minimum_transformer_capacity_kva_at_overload_factor']:.3f}kVA，故选200kVA。"
        "上述容量采用课程短时过载系数，不能直接作为订货参数。",
    )
    add_result_analysis(
        doc,
        "每段设置1套接地源可使分列运行时各母线都有明确的零序回路；母联合闸时只保留健康侧1套，"
        "既维持接地故障电流目标，又避免两套电阻并联。最终仍需由接地变/NGR厂家统一校核零序阻抗、损耗、温升、CT饱和和保护配合。",
    )
    add_note_box(
        doc,
        "防雷与接地边界",
        "MOA绝缘配合不等于直击雷保护完成。缺少雷暴日、建筑高度、土壤电阻率和接地网几何时，不虚构"
        "避雷针保护范围、接地网电阻、接触电压或跨步电压。",
    )

    add_heading(doc, "4.5 母线、绝缘子与套管", 2, page_break_before=False)
    add_body(
        doc,
        "母线型式已从原则选择推进到课程级截面和性能预筛。220kV管母完成载流、热稳、简化弯曲和电晕；"
        "35/10kV矩形母线只完成载流与热稳，电动力、支撑机械、连接金具和厂家温升仍待复核。",
    )
    add_body(
        doc,
        "母线热稳定采用铝导体课程常数C=87和等值时间1.10s。先由几何尺寸求截面积，再计算允许短时电流，"
        "并与本电压级条件性最大短路职责比较。长期载流量另按参考载流量、环境温度和布置方式校核。",
    )
    add_equation(doc, "Ith,allow = C S / (1000 √t)", "4-7")
    add_equation(doc, "Smin = 1000 Ik √t / C", "4-8")
    add_formula_terms(doc, "S为单相导体总截面积，单位mm²；Ith,allow和Ik以kA计，t=1.10s。")
    add_table(
        doc,
        "表4-7 母线课程级选择",
        ["电压级", "课程选择", "持续/热稳定", "补充校验"],
        [
            ["220kV", "铝合金管母Φ100/90mm", "1605.546A＞735.559A；123.785kA＞7.385kA", "弯曲应力10.048＜70MPa；电晕330.444＞145.492kV"],
            ["35kV", "每相3×125×10mm矩形铝母线", "4194A＞3117.691A；311.067kA＞16.291kA", "机械/电动力和厂家温升待核"],
            ["10kV", "每相2×125×10mm矩形铝母线", "3282A＞1909.586A；207.378kA＞15.638kA", "机械/电动力和厂家温升待核"],
        ],
        [2.2, 4.0, 5.5, 3.8],
        8.1,
    )
    add_body(
        doc,
        f"220kV Φ100/90mm管母截面积为{busbars['BUS-220-TUBE-100-90']['calculated']['cross_section_area_mm2']:.3f}mm²，"
        f"41℃修正后载流量{busbars['BUS-220-TUBE-100-90']['calculated']['corrected_ampacity_a']:.3f}A，"
        f"1.10s热稳定允许电流{busbars['BUS-220-TUBE-100-90']['calculated']['thermal_allowable_current_ka']:.3f}kA。"
        f"两者分别高于735.559A持续职责和7.385kA短路职责。",
    )
    add_body(
        doc,
        f"简化电动力计算按3.0m相距、8.0m支撑跨距和1.8动载系数进行，得到弯曲应力"
        f"{busbars['BUS-220-TUBE-100-90']['calculated']['dynamic']['calculated_bending_stress_mpa']:.3f}MPa，低于70MPa课程允许值；"
        f"Peek公式得到电晕起始相电压{busbars['BUS-220-TUBE-100-90']['calculated']['corona']['critical_disruptive_phase_voltage_kv']:.3f}kV，"
        f"高于最高运行相电压{busbars['BUS-220-TUBE-100-90']['calculated']['corona']['highest_operating_phase_voltage_kv']:.3f}kV。",
    )
    add_result_analysis(
        doc,
        "220kV管母已形成载流、热稳、简化机械和电晕的完整课程链。35/10kV矩形母线虽在载流和热稳定上裕度充足，"
        "但缺少支撑跨距、相间距、固有频率、连接受力和厂家温升资料，因此只能写为“已完成电气与热稳定预筛，机械动稳定待核”。",
    )
    add_table(
        doc,
        "表4-8 绝缘子与套管课程目标",
        ["电压级", "最高电压/LIWV", "持续/短时/峰值课程目标", "尚需厂家数据"],
        [
            ["220kV", "252/950kV", "630A；50kA/3s；125kA", "d级；爬距、机械破坏负荷、精确型号"],
            ["35kV", "40.5/185kV", "4000A；31.5kA/4s；80kA", "户外接口d级；爬距、机械与安装图"],
            ["10kV", "12/75kV", "2500A；31.5kA/4s；80kA", "户外接口d级；爬距、机械与安装图"],
        ],
        [2.2, 3.3, 5.2, 4.8],
        8.3,
        "当前35/10kV方案无裸母线直接穿墙，独立裸母线穿墙套管不单列；方案改变时重新选择。",
    )

    add_heading(doc, "第五章 配电装置设计与运行", 1, page_break_before=False)
    add_heading(doc, "5.1 配电装置型式", 2)
    add_table(
        doc,
        "表5-1 配电装置型式选择",
        ["电压等级", "采用型式", "选择理由"],
        [
            ["220kV", "户外AIS、分相中型", "回路不多、课程断面表达清晰、扩建方便、投资低于GIS"],
            ["35kV", "室内金属铠装移开式开关柜", "回路多、分段明确、运行维护方便"],
            ["10kV", "室内金属铠装移开式开关柜", "适合T10进线、SVG和辅助负荷分柜组织"],
            ["0.4kV", "室内低压开关柜", "单母线分段、暗备用和联锁切换"],
        ],
        [2.5, 5.2, 7.8],
        9.4,
    )
    add_body(
        doc,
        "220kV侧本期回路数量较少，户外AIS能够把母线、隔离开关、断路器、CT/CVT、MOA和构架的设备链清晰展开，"
        "便于检修隔离、远期L3扩建以及课程平面与断面表达。GIS虽占地较小、环境适应性好，但设备集成度和投资较高，"
        "且任务书未给出必须采用紧凑布置的场址约束，因此本阶段不选。",
    )
    add_body(
        doc,
        "35kV和10kV回路多、分段与馈线组织明确，采用室内金属铠装移开式开关柜可缩短户外导体长度并便于防误联锁。"
        "0.4kV低压柜则围绕两进线一母联的暗备用方式配置，正常只有一台所用变进线带电。",
    )

    add_heading(doc, "5.2 220kV户外AIS课程几何", 2)
    add_table(
        doc,
        "表5-2 平面与断面主要参数",
        ["参数", "课程采用值", "说明"],
        [
            ["站区范围", "145m×90m", "课程总体边界"],
            ["道路宽度", "4m", "设备运输与检修通道"],
            ["相间距", "4m", "典型三相设备中心距"],
            ["典型功能中心距", "14m", "课程间隔节距"],
            ["母线标高", "9.0m", "断面课程值"],
            ["出线构架标高", "14.5m", "向西北出线"],
            ["主变构架标高", "11.5m", "主变进线构架"],
            ["母线至出线构架", "29m", "平断面一致"],
            ["母线至主变构架/中心", "20m/27m", "平断面一致"],
        ],
        [5.0, 3.5, 7.0],
        9.4,
    )
    add_table(
        doc,
        "表5-3 220kV屋外配电装置课程最小净距",
        ["类别", "采用值/mm", "含义"],
        [
            ["A1", "1800", "带电部分至接地部分"],
            ["A2", "2000", "不同相带电部分之间"],
            ["B1", "2550", "带电部分至栅栏/运输轮廓等课程控制距离"],
            ["B2", "1900", "带电部分至网状遮栏"],
            ["C", "4300", "无遮栏裸导体至地面"],
            ["D", "3800", "检修或不同时间带电部分的课程控制距离"],
        ],
        [2.6, 3.0, 9.9],
        9.4,
        "取自课程讲义表7-2的220J列；施工设计仍须按DL/T 5352-2018全文复核。",
    )
    add_body(
        doc,
        "平面与断面图共用145m×90m站区、4m道路、4m相间距和14m典型功能中心距。母线、线路构架及主变构架标高分别为9.0m、"
        "14.5m和11.5m，使电气净距、设备检修通道和导线引接关系能够在两类图纸中相互核对。课程几何不等于厂家基础图，"
        "但所有尺寸均有统一数据源，避免平面图、断面图各自估算。",
    )
    add_result_analysis(
        doc,
        "A1、A2、B1、B2、C、D净距分别按1800、2000、2550、1900、4300和3800mm控制课程底图。"
        "这些值用于建立可手绘临摹、可复核的布置基线；施工设计阶段仍须结合现行标准全文、海拔、设备外形、端子方向、"
        "运输荷载和基础尺寸重新校核。",
    )
    add_heading(doc, "5.3 事故运行与联锁", 2)
    operation_paragraph = add_body(
        doc,
        "主变N-1时先切除并隔离故障主变，故障清除后闭合35kV母联，并把存活主变负荷限制在180MVA以内，"
        "按当前计算限发35.780%。一回220kV线路退出时，剩余线路也必须受LGJ-400/50温度修正载流量约束。"
        "T10或所用变切换均先隔离故障进线，再投入备用电源，严禁两个健康电源通过低压母联并列。"
        "35/10kV母联合闸前还须退出受电侧接地源，仅保留健康电源侧1套；禁止母联闭合且两套低电阻接地源并联。",
    )
    operation_paragraph.paragraph_format.keep_together = True
    add_body(
        doc,
        "一回220kV线路退出时，剩余线路不能直接承担280.286MVA全部出力，因为41℃课程允许载流量只有475.242A。"
        "执行主变N-1限发后线路电流为472.377A，虽然满足当前课程参考，但仅余2.864A，因此事故运行时应同时监视主变容量和线路温度。",
    )
    add_body(
        doc,
        "T10事故转供除检查31.5MVA容量外，还应限制两套SVG在吸收无功方向的最严组合；课程计算得到单台T10最大负载率79.17%。"
        "所用变切换则遵循“先断故障进线、再确认母联和备用进线闭锁、最后投入备用”的顺序，防止两台低压变压器并列。",
    )
    add_figure(
        doc,
        ROOT / "drawings" / "exports" / "switchyard_plan_a1.png",
        "图5-1 LAY-220-01 220kV户外AIS配电装置平面布置图（A1，1:200，课程设计底图）",
    )
    add_figure(
        doc,
        ROOT / "drawings" / "exports" / "switchyard_section_a1.png",
        "图5-2 SEC-220-01 220kV户外AIS典型间隔断面图（A1，1:100，课程设计底图）",
    )
    add_optional_figure(
        doc,
        ROOT / "drawings" / "exports" / "switchyard_line_bay_detail_a1.png",
        "图5-3 SEC-220-L1-01 220kV I段L1线路间隔断面详图（A1，1:50，课程设计增强详图）",
        missing_message=(
            "预期文件drawings/exports/switchyard_line_bay_detail_a1.png尚不存在。正式再生成前必须完成第四图"
            "导出与目检，不得把本提示版作为最终提交版。"
        ),
    )

    add_heading(doc, "第六章 结论与设计边界", 1)
    add_heading(doc, "6.1 主要设计结论", 2)
    add_body(
        doc,
        "本设计形成了从主接线、负荷、短路到设备和布置的同源闭环。各项结论并非孤立的表格结果："
        "主接线确定运行方式，运行方式决定持续电流与短路职责，职责再控制设备等级、母线截面和事故操作边界。",
    )
    add_bullets(
        doc,
        [
            "主接线采用220kV、35kV和10kV单母线分段；220kV分段断路器、35/10kV母联正常断开，0.4kV母联正常闭合。",
            "主变采用2×180MVA、220/35kV、YNd11、uk=14%，正常负载率77.857%；N-1需限发35.780%。",
            "采用2×31.5MVA、35/10.5kV、YNd11、uk=8%的T10和2×±12Mvar SVG，补偿后功率因数约0.98147。",
            "所用变采用2×200kVA、SCB14、Dyn11、uk=4%，暗备用，正常一台运行。",
            "10kV条件性最大课程短路电流15.638kA，设备按31.5kA开断等级预筛。",
            "220kV配电装置采用户外AIS分相中型，35/10kV采用室内金属铠装移开式开关柜。",
            "35kV采用2×1000kVA、10kV采用2×200kVA ZN接地变+低电阻；正常每段1套，母联转供时仅保留健康侧1套。",
            "CAD按三张必交图加SEC-220-L1-01增强详图组织，共四张A1图。",
        ],
    )
    add_result_analysis(
        doc,
        "本方案的主要约束是主变N-1限发35.780%、LGJ-400/50高温事故载流裕度很小、35kV 3150A柜持续电流裕度约1.04%，"
        "以及35kV MOA持续电压裕度仅约0.74%。这些项目均已在计算书中保留完整计算链，并在最终工程边界中列为优先复核项。",
    )
    add_heading(doc, "6.2 课程假设与待复核项目", 2)
    add_body(
        doc,
        "课程设计必须在信息不完备时形成可检验方案，但不能把假设写成施工结论。下表按“当前已完成的课程闭环”和"
        "“仍需真实场址、标准全文、保护专题或厂家资料关闭的项目”进行分层，便于后续替换输入时同步重算、重绘和重写。",
    )
    add_table(
        doc,
        "表6-1 课程冻结值与最终工程边界",
        ["项目", "课程设计采用", "最终工程复核"],
        [
            ["短路方法", "100MVA标幺X-only；k=1.8；变流器1.1~1.2倍", "按GB/T 15544.1-2023完整R/X、电压系数、κ和开断时刻"],
            ["设备", "额定值等级预筛", "精确厂家型号、41℃修正、外绝缘与型式试验"],
            ["无功", "目标pf=0.98；2×±12Mvar", "无功电压、谐波与SVG动态专题"],
            ["CT/PT/ZCT", "三相CT/PT目标及35/10kV电缆回路零序CT", "负担、饱和、暂态性能、窗口尺寸、二次容量和电缆压降"],
            ["导体与母线", "220kV载流/热稳/简化弯曲/电晕；35/10kV载流/热稳", "35/10kV电动力与支撑机械、厂家温升、共振和连接金具"],
            ["MOA", "Ur/Uc/In/残压与LIWV裕度课程预筛", "接地方式、10s TOV、能量、引线压降和厂家型号"],
            ["场址与布置", "海拔≤1000m、污秽d级、课程几何", "真实场址、DL/T 5352全文和厂家外形基础图"],
            ["接地与直击雷", "接地设备包、母联切换顺序和设备接地接口", "单相接地/零序、厂家联锁、土壤电阻率、接触/跨步电压和避雷针范围"],
        ],
        [3.2, 5.2, 7.1],
        8.7,
    )

    add_heading(doc, "6.3 图纸和计算成果清单", 2)
    add_table(
        doc,
        "表6-2 本阶段成果",
        ["编号", "成果", "文件"],
        [
            ["SLD-01", "电气主接线简图", "drawings/exports/\nsingle_line_a1.pdf"],
            ["LAY-220-01", "220kV户外AIS平面布置图", "drawings/exports/\nswitchyard_plan_a1.pdf"],
            ["SEC-220-01", "220kV户外AIS典型间隔断面图", "drawings/exports/\nswitchyard_section_a1.pdf"],
            ["SEC-220-L1-01", "220kV I段L1线路间隔断面详图", "drawings/exports/\nswitchyard_line_bay_detail_a1.pdf"],
            ["CALC-LT", "负荷、变压器与持续电流结果", "calculations/results/\nload_and_transformer_summary.md"],
            ["CALC-SC", "短路计算结果", "calculations/results/short_circuit/\nshort_circuit_summary.md"],
            ["CALC-EQ", "设备额定值预筛", "calculations/results/equipment_selection/\nequipment_selection_summary.md"],
        ],
        [2.7, 5.1, 7.7],
        8.5,
        margin_y=55,
    )
    add_references(doc, REFERENCES, page_break_before=False)

    path = REPORT_DIR / "01_220kV新能源汇集变电所_技术设计说明书.docx"
    finalize_document(doc, path)
    return path


def build_calculation_book(data: dict) -> Path:
    load = data["load"]
    sc = data["sc"]
    equipment = data["equipment"]
    load35 = load["load_35kv"]
    load10 = load["load_10kv"]
    main_transformer = load["main_transformer"]
    reactive = load["reactive_compensation"]
    station_service = load["station_service"]
    outgoing = load["outgoing_220kv"]
    outgoing_ampacity = load["outgoing_220kv_conductor_ampacity"]
    sc_points = sc["points"]
    network = sc["network"]
    course_completion = equipment["course_completion"]
    busbars = {item["id"]: item for item in course_completion["busbars"]}
    arresters = {item["id"]: item for item in course_completion["surge_arresters"]}
    grounding_packages = {
        item["id"]: item
        for item in course_completion["grounding_transformer_resistor_packages"]
    }
    doc = new_document("技术设计计算书")

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_front_section(front)
    add_toc(doc)

    main_sec = start_main_section(doc, SHORT_TITLE + "—技术设计计算书")

    add_heading(doc, "计算说明与统一基准", 1, page_break_before=False)
    add_body(
        doc,
        "本计算书保留负荷、容量、持续电流、标幺电抗、短路电流和设备校验的公式、代入和中间结果。"
        "除特别说明外，功率单位采用MW/Mvar/MVA，电压采用kV，电流采用A或kA。短路计算为任务书"
        "X-only课程初算，最终工程值仍须按现行标准和厂家模型复核。",
    )
    add_body(
        doc,
        "计算遵循“原始数据—通用公式—参数代入—中间结果—判据比较—结论”的顺序。"
        "表格用于汇总同类对象，关键选择不直接从表格读取结论，而是在正文中给出至少一个完整代入示例和控制因素分析。"
        "所有未取整数值由计算脚本保存，正文通常保留3~6位小数，设备等级按额定标准值向上选择。",
    )
    add_body(
        doc,
        "标量相加、P-Q向量合成、正常分段、事故转供和禁止并列属于不同计算口径。"
        "本书在每次使用这些口径时均说明适用对象；禁止方式仅用于敏感性，不自动转化为日常设备必选职责。"
        "课程假设值在式中正常参与计算，但最终结论仍保留被教师口径、标准全文或厂家数据覆盖的接口。",
    )
    add_table(
        doc,
        "表0-1 统一计算条件",
        ["项目", "采用值", "说明"],
        [
            ["短路基准容量SB", "100MVA", "全网统一"],
            ["平均额定电压UB", "230/37/10.5kV", "220/35/10kV级"],
            ["35kV同时出力系数", "0.95", "仅用于主变综合容量"],
            ["10kV同时系数", "0.80", "辅助负荷综合计算"],
            ["线损率", "5%", "主结果按乘1.05"],
            ["最高环境温度", "41℃", "导体和设备环境校核"],
            ["峰值系数k", "1.8", "课程假设"],
            ["热稳定等值时间", "1.10s", "后备保护1.00s+总开断0.08s向上取整"],
            ["变流器短路贡献", "1.1~1.2倍额定电流", "算术相加课程上界"],
        ],
        [5.0, 4.0, 6.5],
        9.3,
    )
    add_heading(doc, "主要符号", 2)
    symbol_paragraph = add_body(
        doc,
        "P、Q、S分别表示有功、无功和视在功率；cosφ为功率因数；Imax为最大持续工作电流；"
        "X*为标幺电抗；Ik为三相短路初始对称电流；ip为课程峰值电流；Ith和tr分别为设备短时耐受电流及持续时间。",
    )
    symbol_paragraph.paragraph_format.keep_together = True
    add_body(
        doc,
        "下标B表示基准量，N表示额定量，Σ表示故障点至电源的总等值量；normal表示正常运行，N-1表示任一关键元件退出，"
        "closed表示分段或母联条件闭合。式中电压均取线电压，除接地电阻计算明确使用相电压Uφ外，不混用相、线电压。",
    )

    add_heading(doc, "第一章 负荷及变压器容量计算", 1)
    add_heading(doc, "1.1 35kV新能源负荷", 2)
    add_body(
        doc,
        "35kV负荷由风电场A/B、光伏电站A/B和储能电站组成。容量计算先把各项有功功率按自身功率因数折算为视在功率和无功功率；"
        "主变综合容量再统一施加0.95最大同时出力系数和5%线损。单回馈线电流另按该类最大出力除以实际回路数计算，不施加0.95系数。",
    )
    add_equation(doc, "Sᵢ = Pᵢ / cosφᵢ", "1-1")
    add_equation(doc, "Qᵢ = Pᵢ tan(arccos(cosφᵢ))", "1-2")
    add_formula_terms(
        doc,
        "Pᵢ以MW计，Sᵢ以MVA计，Qᵢ以Mvar计；单回设计电流Iᵢ=Pᵢ(1+0.05)/(n√3Ucosφᵢ)，其中n为回路数。",
    )
    add_body(doc, "以风电场A为例：")
    add_equation(doc, "SWA = 80 / 0.95 = 84.210526 MVA", compact=True)
    add_equation(doc, "QWA = 80 × tan(arccos0.95) = 26.294728 Mvar", compact=True)
    add_equation(doc, "IWA = 84.210526 × 1.05 / (4 × √3 × 35) × 1000 = 364.642 A", compact=True)
    add_body(
        doc,
        "风电场B为60MW、3回，单回有功同样为20MW，因此与风电场A得到相同的单回电流364.642A。"
        "光伏电站A功率因数较高但每回承担30MW，单回电流增至530.220A；储能只有1回线路，形成全站最大35kV单回馈线电流546.963A。",
    )
    rows = []
    for item in load["load_35kv"]["items"]:
        rows.append(
            [
                item["label_zh"],
                fmt(item["active_power_mw"], 1),
                fmt(item["power_factor"], 2),
                item["circuits"],
                fmt(item["base_apparent_mva"], 3),
                fmt(item["base_reactive_mvar"], 3),
                fmt(item["per_circuit_current_a"], 3),
            ]
        )
    add_table(
        doc,
        "表1-1 35kV负荷逐项计算",
        ["负荷", "P/MW", "cosφ", "回路", "S/MVA", "Q/Mvar", "单回设计电流/A"],
        rows,
        [2.6, 1.9, 1.8, 1.7, 2.0, 2.1, 3.4],
        8.0,
        "单回设计电流已计5%线损，不乘0.95同时出力系数。",
        nowrap_header_columns={1, 2, 3, 4, 5},
        nowrap_body_columns={1, 2, 3, 4, 5, 6},
        margin_x=75,
    )
    add_equation(doc, "ΣS = 280.988185 MVA", "1-3")
    add_equation(doc, "S35 = 280.988185 × 0.95 × 1.05 = 280.285714 MVA", "1-4")
    add_body(
        doc,
        f"逐项无功合计为{load35['gross_reactive_mvar']:.6f}Mvar。若按P、Q向量合成，计同时系数和线损后的有功、无功分别为"
        f"{load35['with_losses_active_mw']:.3f}MW和{load35['with_losses_reactive_mvar']:.6f}Mvar，故：",
    )
    add_equation(
        doc,
        f"S35,vector = √({load35['with_losses_active_mw']:.3f}² + "
        f"{load35['with_losses_reactive_mvar']:.6f}²) = {load35['vector_with_losses_mva']:.6f} MVA",
        compact=True,
    )
    add_body(
        doc,
        "向量法按P、Q分别计同时系数和线损得到279.840482MVA；任务书课程主口径采用逐项视在功率和的"
        "标量处理，得到280.285714MVA。两者接近，正式容量校验取后者。",
    )
    add_result_analysis(
        doc,
        f"标量口径比向量口径高{load35['with_losses_mva'] - load35['vector_with_losses_mva']:.6f}MVA，"
        "保持了课程计算的偏保守性。后续主变容量和220kV送出电流均以280.285714MVA为主结果，"
        "单回馈线电流仍按各类自身最大值计算，以免同时系数被重复使用。",
    )

    add_heading(doc, "1.2 10kV辅助负荷", 2)
    add_body(
        doc,
        "10kV负荷包括站用电备用电源、无功补偿及冷却系统、集控通信及监控系统。综合容量计算采用0.80同时系数和5%线损；"
        "各馈线设备电流仍按本回路额定最大负荷计5%线损，不乘0.80，以保证单回路选型不因全站同时系数而偏小。",
    )
    add_equation(doc, "ΣS10 = 0.8/0.85 + 0.6/0.85 + 0.4/0.90 = 2.091503 MVA", compact=True)
    add_equation(doc, "Istation-backup = 0.8×1.05/(√3×10×0.85)×1000 = 57.056 A", compact=True)
    rows = []
    for item in load["load_10kv"]["items"]:
        rows.append(
            [
                item["label_zh"],
                fmt(item["active_power_mw"], 1),
                fmt(item["power_factor"], 2),
                fmt(item["base_apparent_mva"], 3),
                fmt(item["per_circuit_current_a"], 3),
            ]
        )
    add_table(
        doc,
        "表1-2 10kV辅助负荷逐项计算",
        ["负荷", "P/MW", "cosφ", "S/MVA", "单回设计电流/A"],
        rows,
        [5.0, 2.0, 2.0, 2.7, 3.8],
        9.2,
    )
    add_equation(doc, "S10 = 2.091503 × 0.80 × 1.05 = 1.756863 MVA", "1-5")
    add_result_analysis(
        doc,
        "1.756863MVA用于10kV辅助负荷综合分析；57.056A、42.792A和26.943A分别用于三类单回馈线额定电流选择。"
        "两种结果的用途不同，不能用101.433A的综合母线电流替代任一具体馈线，也不能用单回最大值简单推回主变综合容量。",
    )

    add_heading(doc, "1.3 主变压器容量与N-1", 2)
    add_body(
        doc,
        "主变容量按35kV新能源汇集主口径280.285714MVA校验。正常方式下两台主变均衡分担；N-1方式假定任一台退出，"
        "另一台不允许超过180MVA额定容量，超出部分由新能源限发消除。该策略适用于可控新能源汇集站，但必须作为运行约束写入说明书。",
    )
    add_equation(doc, "Snormal,each = 280.285714 / 2 = 140.142857 MVA", "1-6")
    add_equation(doc, "βnormal = 140.142857 / 180 × 100% = 77.857%", "1-7")
    add_equation(doc, "KN-1 = 180 / 280.285714 × 100% = 64.220%", "1-8")
    add_equation(doc, "ΔS = 280.285714 − 180 = 100.285714 MVA", "1-9")
    add_equation(doc, "Kcurtail = (1 − 180 / 280.285714) × 100% = 35.7798%", "1-10")
    add_body(
        doc,
        "主变课程型式冻结为2×180MVA、220/35kV双绕组、YNd11、uk=14%；220kV星形中性点经专用CT"
        "直接接地，35kV侧为三角形绕组。",
    )
    add_equation(doc, "IT,220 = 180/(√3×220) = 472.377 A", compact=True)
    add_equation(doc, "Imax,T,220 = 1.05×472.377 = 495.996 A", compact=True)
    add_equation(doc, "IT,35 = 180/(√3×35) = 2969.230 A", compact=True)
    add_equation(doc, "Imax,T,35 = 1.05×2969.230 = 3117.691 A", compact=True)
    add_body(
        doc,
        "若保守叠加10kV辅助负荷，所需容量为282.042577MVA，正常负载率78.345%，N-1缺口"
        "102.042577MVA。该值作为敏感性，不改变主校核口径。",
    )
    add_result_analysis(
        doc,
        "正常负载率77.857%说明2×180MVA容量配置在日常方式下具有合理裕度；N-1覆盖率64.220%则表明该裕度不足以无条件承担全站最大出力。"
        "1.05系数只用于回路设备额定电流选择，不能被解释为主变本体允许长期过载5%。",
    )

    add_heading(doc, "1.4 所用变压器", 2)
    add_body(
        doc,
        "所用电清单先按项目电压等级筛选，110kV配电装置电源和110kV断路器加热两项因本站无110kV系统而剔除。"
        "其余负荷按连续、经常短时和不经常短时分类；空气压缩机总装机44kW，基准场景按一台22kW同时投入，最严经常场景按44kW全部投入。",
    )
    add_table(
        doc,
        "表1-3 所用电负荷统计结果",
        ["统计口径", "功率/kW"],
        [
            ["连续适用负荷", "96.3"],
            ["短时适用负荷", "89.0"],
            ["明确经常负荷", "81.2"],
            ["明确不经常负荷", "68.1"],
            ["未标频度连续负荷", "36.0"],
            ["经常短时负荷", "44.0"],
            ["短时同时投入值", "22.0"],
        ],
        [9.5, 6.0],
        9.2,
        margin_y=45,
    )
    add_body(
        doc,
        "连续负荷96.3kW包括主变风机、充电及通风、照明和冬季加热等项目。基准场景Pbase=96.3+22.0=118.3kW；"
        "最严经常场景Pworst=96.3+44.0=140.3kW。容量计算统一取功率因数0.80并留10%裕度。",
    )
    add_equation(doc, "ST = Pscenario(1+Km)/cosφ", compact=True)
    add_formula_terms(doc, "Km=0.10，cosφ=0.80；标准容量从100、125、160、200、250、315kVA序列中向上选取。")
    add_equation(doc, "Sbase = 118.3 × 1.10 / 0.80 = 162.6625 kVA", "1-11")
    add_equation(doc, "Sworst = 140.3 × 1.10 / 0.80 = 192.9125 kVA", "1-12")
    add_body(doc, "选用2×200kVA、10/0.4kV、SCB14干式、Dyn11、uk=4%的所用变，暗备用；正常一台运行。")
    add_result_analysis(
        doc,
        "192.9125kVA决定标准容量取200kVA。两台设备并非各承担一半负荷，而是每台都应能独立承担最严课程场景；"
        "正常一台运行、0.4kV母联闭合，备用投入时必须防止两进线并列。电动机启动压降和精确负荷功率因数仍需在最终设备资料齐全后复核。",
    )

    add_heading(doc, "1.5 SVG与T10", 2)
    add_equation(doc, "Qc = P [tanφ1 − tanφ2]", "1-13")
    add_body(
        doc,
        "T10冻结为2×31.5MVA、35/10.5kV、YNd11、uk=8%，两段各设±12Mvar SVG。补偿计算不抵扣新能源变流器自身无功裕度，"
        "先对35kV汇集负荷计算，再保守叠加10kV辅助负荷。",
    )
    add_equation(
        doc,
        "Qc,35 = 75.991709 − 269.325×tan(arccos0.98) = 21.302935 Mvar",
        "1-14",
    )
    add_equation(
        doc,
        "Qc,Σ = 76.883260 − 270.837×tan(arccos0.98) = 21.887462 Mvar",
        "1-15",
    )
    add_equation(doc, "Kmargin = (24/21.887462 − 1)×100% = 9.65%", "1-16")
    add_equation(
        doc,
        "cosφ' = 270.837 / √[270.837² + (76.883260−24)²] = 0.98147",
        "1-17",
    )
    add_body(
        doc,
        "T10容量按两套SVG的无功交换范围和10kV辅助有功共同校验。正常均衡时每台承担一半辅助有功、辅助无功和一套12Mvar SVG；"
        "N-1最严吸收工况由单台T10承担全部辅助负荷和两套SVG吸收无功。",
    )
    add_equation(
        doc,
        "ST10,normal = √[(1.512/2)² + (12+0.891552/2)²] = 12.468716 MVA，β=39.583%",
        "1-18",
    )
    add_equation(
        doc,
        "ST10,N-1 = √[1.512² + (24+0.891552)²] = 24.937431 MVA，β=79.166%",
        "1-19",
    )
    add_equation(doc, "IT10,35 = 31.5 / (√3 × 35) = 519.615 A", "1-20")
    add_equation(doc, "1.05 IT10,35 = 545.596 A", "1-21")
    add_equation(doc, "IT10,10.5 = 31.5 / (√3 × 10.5) = 1732.051 A", "1-22")
    add_equation(doc, "1.05 IT10,10.5 = 1818.653 A", "1-23")
    add_equation(doc, "1.05 IT10,10kV设备口径 = 31.5/(√3×10)×1.05 = 1909.586 A", "1-24")
    add_result_analysis(
        doc,
        "2×±12Mvar使补偿后功率因数达到0.98147，且单台T10在最严N-1吸收工况下负载率仍低于80%。"
        "式(1-23)按10.5kV绕组电压计算，式(1-24)按10kV开关设备校核电压计算，二者对象不同；"
        "最终仍须由无功电压、谐波和SVG厂家持续过载曲线覆盖。",
    )
    add_heading(doc, "第二章 各回路最大持续工作电流计算", 1)
    add_heading(doc, "2.1 通用公式与主变回路", 2)
    add_body(
        doc,
        "持续电流必须针对具体回路和运行方式确定。主变回路按设备额定容量乘1.05选型裕度；220kV线路正常按两回均分，"
        "N-1按单回承担全部本期送出；母联按可能转供的最大设备容量确定。1.05只用于电器及导体选型，不改变变压器额定容量。",
    )
    add_equation(doc, "Imax = Smax / (√3 Un)", "2-1")
    add_equation(doc, "IT,220 = 180 / (√3 × 220) × 1.05 = 495.996 A", "2-2")
    add_equation(doc, "IT,35 = 180 / (√3 × 35) × 1.05 = 3117.691 A", "2-3")
    add_equation(doc, "IL,normal = 280.285714 / (2√3 × 220) = 367.780 A", "2-4")
    add_equation(doc, "IL,N-1 = 280.285714 / (√3 × 220) = 735.559 A", "2-5")
    add_result_analysis(
        doc,
        "220kV线路与分段回路的持续职责取735.559A，以覆盖一回线路退出或一段事故转供。主变高、低压侧分别取495.996A和3117.691A。"
        "这些值来自不同事故路径，不能用正常双回每回367.780A替代主变或母联设备职责。",
    )

    add_heading(doc, "2.2 35kV馈线与母线段", 2)
    feeder_rows = []
    for item in load["load_35kv"]["items"]:
        feeder_rows.append(
            [
                item["label_zh"],
                item["circuits"],
                fmt(item["base_per_circuit_mva"], 3),
                fmt(item["per_circuit_design_mva"], 3),
                fmt(item["per_circuit_current_a"], 3),
            ]
        )
    add_table(
        doc,
        "表2-1 35kV馈线持续电流",
        ["负荷", "回路数", "单回基础S/MVA", "计损耗S/MVA", "Imax/A"],
        feeder_rows,
        [3.5, 2.0, 3.0, 3.0, 4.0],
        9.0,
        margin_y=45,
    )
    add_equation(doc, "I35-I,base = 2230.732 A；I35-II,base = 2404.371 A", "2-6")
    add_body(
        doc,
        "式(2-6)按各馈线未计线损的基础电流和实际分配累加，不把全站4623.515A简单平均。"
        "为避免与表2-1的计损耗单回电流混淆，还需并列说明另外两种口径。",
    )
    add_equation(doc, "I35-I,group = 2225.155 A；I35-II,group = 2398.360 A", compact=True)
    add_body(
        doc,
        "上式按各馈线基础视在功率标量求和后统一乘0.95同时系数和1.05线损，适合与全站课程主口径对应。"
        "若按各段P、Q分别合成，则I、II段分别为2221.564A和2394.612A；若把每回馈线都按自身最大出力计5%线损后直接相加，"
        "I、II段分别为2342.268A和2524.590A，表示各馈线同时达到各自最大值的保守和。",
    )
    add_equation(doc, "I35-I,vector = 2221.564 A；I35-II,vector = 2394.612 A", compact=True)
    add_equation(doc, "I35-I,max-sum = 2342.268 A；I35-II,max-sum = 2524.590 A", compact=True)
    add_result_analysis(
        doc,
        "2230.732/2404.371A、2225.155/2398.360A、2221.564/2394.612A和2342.268/2524.590A分别对应基础分配、"
        "标量综合、P-Q向量综合和逐馈线最大值四种口径。"
        "35kV主变进线、母联及主母线设备最终仍按3117.691A事故转供职责选型，不采用其中任何较小值。",
    )

    add_heading(doc, "2.3 10kV及SVG回路", 2)
    add_body(
        doc,
        "10kV进线和母联按T10额定容量在10kV开关设备电压基准下计算，SVG按12Mvar、10.5kV额定端电流计算，"
        "其他馈线按各自有功、功率因数和5%线损计算。下表的最大值服务于单个回路选型，不代表10kV母线综合负荷。",
    )
    add_table(
        doc,
        "表2-2 10kV主要回路持续电流",
        ["回路", "Imax/A", "计算依据"],
        [
            ["T10进线/10kV母联", "1909.586", "31.5MVA按10kV设备口径×1.05"],
            ["单套SVG", "692.820", "12Mvar按10.5kV额定电流×1.05"],
            ["站用电备用电源", "57.056", "0.8MW、pf=0.85、计5%线损"],
            ["无功补偿及冷却", "42.792", "0.6MW、pf=0.85、计5%线损"],
            ["集控通信及监控", "26.943", "0.4MW、pf=0.90、计5%线损"],
        ],
        [5.0, 3.5, 7.0],
        9.4,
    )
    add_equation(doc, "ISVG = 12/(√3×10.5)×1.05×1000 = 692.820 A", compact=True)
    add_equation(doc, "Ibackup = 0.8×1.05/(√3×10×0.85)×1000 = 57.056 A", compact=True)
    add_result_analysis(
        doc,
        "10kV进线/母联1909.586A控制开关柜额定电流，SVG 692.820A控制SVG馈线柜；其余辅助馈线均远低于1250A等级。"
        "进线选择2500A后保留590.414A裕度，同时避免把按10.5kV绕组计算的1818.653A直接用于10kV柜口径。",
    )

    add_heading(doc, "2.4 LGJ-400/50温度修正", 2)
    add_body(
        doc,
        "导线载流量随环境温度升高而下降。课程采用参考环境25℃、允许导体温度70℃和参考载流量592A，"
        "按平方根温度修正公式换算至本站最高41℃。该方法未包含风速、日照、导体表面状态和动态热额定，因此只作课程保守校核。",
    )
    add_equation(doc, "KT = √[(70−41)/(70−25)] = 0.802773", "2-7")
    add_equation(doc, "Iallow,41 = 592 × 0.802773 = 475.242 A", "2-8")
    add_equation(doc, "Kline-curtail = (1−475.242/735.559)×100% = 35.390%", "2-9")
    add_equation(doc, "Imain-T-N-1 = 180/(√3×220) = 472.377 A", "2-10")
    line_ampacity_paragraph = add_body(
        doc,
        "主变N-1后的线路电流472.377A，小于课程允许值475.242A，余量2.864A。因此主变限发35.780%"
        "略严于线路限发35.390%，但厂家热额定资料到位后必须覆盖复核。",
    )
    line_ampacity_paragraph.paragraph_format.keep_together = True
    add_result_analysis(
        doc,
        "正常双回每回367.780A有107.462A裕度；单回承担全部出力时短缺260.318A，不能无条件运行。"
        "主变N-1限发恰好把线路电流压至472.377A，但2.864A裕度极小，所以事故运行时主变容量和线路热稳定是同时存在的控制条件。",
    )

    add_heading(doc, "第三章 短路电流计算", 1, page_break_before=False)
    add_heading(doc, "3.1 标幺值与等值网络", 2)
    add_body(
        doc,
        "本章采用100MVA统一基准和230/37/10.5kV平均额定电压，电阻在课程初算中忽略。"
        "标幺法先求各电压级基准阻抗、基准电流，再把系统、线路和变压器电抗换算到同一基准。"
        "计算得到的是三相短路初始对称电流，用于设备等级预筛。",
    )
    add_equation(doc, "ZB = UB²/SB；IB = SB/(√3UB)；Ik = IB/XΣ*", compact=True)
    add_equation(doc, "IB,220 = 100/(√3×230) = 0.251022 kA", compact=True)
    add_equation(doc, "IB,35 = 100/(√3×37) = 1.560562 kA", compact=True)
    add_equation(doc, "IB,10 = 100/(√3×10.5) = 5.498574 kA", compact=True)
    add_body(
        doc,
        "系统电抗按自身容量基准换算，线路电抗按xL/ZB换算，变压器电抗按短路电压百分数和额定容量换算。"
        "系统1经L1、系统2经L2构成两条220kV电源路径；分段闭合时两路径并联，向低电压侧继续串入主变和T10电抗。",
    )
    add_equation(doc, "ZB,220 = UB²/SB = 230²/100 = 529 Ω", "3-1")
    add_equation(doc, "X*sys1 = 0.40×100/2400 = 0.016667", "3-2")
    add_equation(doc, "X*sys2 = 0.45×100/2000 = 0.022500", "3-3")
    add_equation(doc, "X*L1 = 0.40×70×100/230² = 0.052930", "3-4")
    add_equation(doc, "X*L2 = 0.40×85×100/230² = 0.064272", "3-5")
    add_equation(doc, "X*path1 = 0.069597；X*path2 = 0.086772", "3-6")
    add_equation(doc, "X*parallel = X1X2/(X1+X2) = 0.038621", "3-7")
    add_equation(doc, "X*T = 0.14×100/180 = 0.077778", "3-8")
    add_equation(doc, "X*T10 = 0.08×100/31.5 = 0.253968", "3-9")
    add_table(
        doc,
        "表3-1 元件标幺电抗",
        ["元件", "X*/pu", "说明"],
        [
            ["系统1", "0.016667", "以2400MVA为原基准"],
            ["L1", "0.052930", "70km，0.40Ω/km"],
            ["系统1+L1", "0.069597", "220kV-I段分列电源路径"],
            ["系统2", "0.022500", "以2000MVA为原基准"],
            ["L2", "0.064272", "85km，0.40Ω/km"],
            ["系统2+L2", "0.086772", "220kV-II段分列电源路径"],
            ["两电源并联", "0.038621", "220kV分段条件闭合"],
            ["每台主变", "0.077778", "180MVA，uk=14%"],
            ["每台T10", "0.253968", "31.5MVA，uk=8%"],
        ],
        [5.0, 3.0, 7.5],
        9.4,
    )
    add_body(
        doc,
        "新能源场站和SVG短路贡献采用额定电流1.1~1.2倍的RMS算术上界。对任一变流器支路，"
        "课程最大贡献按Iconv,max=1.2Sr/(√3U)计算；综合对称电流为电网分量与变流器上界相加。",
    )
    add_equation(doc, "Icourse,max = Ik,grid + ΣIconv,max", compact=True)
    add_equation(doc, "ip,course = √2 × 1.8 × Icourse,max", compact=True)
    add_result_analysis(
        doc,
        "固定k=1.8同时作用于电网和变流器RMS算术上界，只用于形成保守的设备峰值预筛。"
        "完整工程计算还需考虑R/X、电压系数、变压器修正、开断时刻以及厂家变流器限流波形，因此本章不声称实现了GB/T 15544.1-2023全部方法。",
    )

    add_heading(doc, "3.2 220kV短路点", 2)
    add_body(
        doc,
        "220kV分列运行时，I段和II段分别由系统1+L1、系统2+L2供电。分段条件闭合时两路径并联，"
        "等值电抗下降到0.038621，故该场景形成220kV设备预筛的最大电网分量。",
    )
    add_equation(doc, "Ik,220-I = 100/(√3×230×0.069597) = 3.607 kA", "3-10")
    add_body(doc, "计I段新能源贡献上界0.426kA，课程控制职责为4.033kA，课程峰值10.266kA；3.607kA仅为电网分量。")
    add_equation(doc, "I220-I,total = 3.606806 + 0.425867 = 4.032673 kA", compact=True)
    add_equation(doc, "ip,220-I = √2×1.8×4.032673 = 10.265509 kA", compact=True)
    add_equation(doc, "Ik,220-II = 100/(√3×230×0.086772) = 2.893 kA", "3-11")
    add_body(doc, "计II段新能源贡献上界0.459kA，课程综合上限为3.352kA。")
    add_equation(doc, "Ik,220-closed = 100/(√3×230×0.038621) = 6.500 kA", "3-12")
    add_body(
        doc,
        "6.500kA和电网峰值16.546kA仅为电网分量。计全部新能源贡献上界0.885kA后，条件性控制职责"
        "为7.385kA，课程峰值18.798kA；该场景必须以系统允许并列为前提。",
    )
    add_equation(doc, "I220,closed,total = 6.499689 + 0.884883 = 7.384572 kA", compact=True)
    add_equation(doc, "ip,220,closed = √2×1.8×7.384572 = 18.798051 kA", compact=True)
    add_result_analysis(
        doc,
        "正常分列时I段综合4.033kA大于II段3.352kA；条件闭合时升至7.385kA。"
        "设备条件性预校核采用7.385/18.798kA，而运行规程仍把分段闭合作为需调度许可和同期检查的条件方式。",
    )

    add_heading(doc, "3.3 35kV短路点", 2)
    add_body(
        doc,
        "35kV故障路径在220kV电源路径后串入一台主变电抗0.077778。正常分列时两段上游路径不同，"
        "因此I段电网短路电流高于II段；220kV分段闭合后，两段上游等值相同，新能源接入规模较大的II段形成条件性最大综合值。",
    )
    add_equation(doc, "XΣ,35-I = 0.069597 + 0.077778 = 0.147375", "3-13")
    add_equation(doc, "Ik,35-I,grid = 100/(√3×37×0.147375) = 10.588 kA", "3-14")
    add_body(doc, "加I段新能源贡献上界2.677kA，综合上限13.265kA。")
    add_equation(doc, "XΣ,35-II = 0.086772 + 0.077778 = 0.164550", "3-15")
    add_equation(doc, "Ik,35-II,grid = 9.483 kA；综合上限12.368 kA", "3-16")
    add_body(
        doc,
        "220kV分段条件闭合时，两段上游等值相同，35kV-II段因新能源额定电流较大形成最大综合上限"
        "16.291kA，课程峰值41.470kA。两台健康主变低压侧并列时综合敏感性25.694kA，属于禁止方式。",
    )
    add_equation(doc, "I35-II,closed,total = 13.405738 + 2.885245 = 16.290983 kA", compact=True)
    add_equation(doc, "ip,35,max = √2×1.8×16.290983 = 41.470074 kA", compact=True)
    add_equation(doc, "I35,parallel,total = 25.693932 kA；ip = 65.406073 kA（禁止方式）", compact=True)
    add_result_analysis(
        doc,
        "35kV开关设备按16.291kA条件职责预筛，低于31.5kA等级；25.694kA禁止并列值虽仍低于31.5kA，"
        "但不能据此允许两台健康主变长期并列，因为其峰值、保护配合、接地方式和运行许可均发生改变。",
    )

    add_heading(doc, "3.4 10kV短路点", 2)
    add_body(
        doc,
        "10kV故障路径在主变后继续串入T10电抗0.253968。由于T10电抗占总电抗比例较大，"
        "正常分段时I、II段电网短路电流分别为13.700kA和13.138kA；每段再叠加一套SVG的1.2倍额定电流上界0.792kA。",
    )
    add_equation(doc, "XΣ,10-I = 0.069597+0.077778+0.253968 = 0.401343", "3-17")
    add_equation(doc, "Ik,10-I,grid = 13.700 kA；计SVG后14.492 kA", "3-18")
    add_equation(doc, "XΣ,10-II = 0.086772+0.077778+0.253968 = 0.418518", "3-19")
    add_equation(doc, "Ik,10-II,grid = 13.138 kA；计SVG后13.930 kA", "3-20")
    add_body(
        doc,
        "220kV分段条件闭合时，单台T10供本段的综合上限为15.638kA，课程峰值39.808kA。两台健康T10"
        "经10kV母联并列时综合敏感性为28.472kA，仅用于联锁风险提示。",
    )
    add_equation(doc, "I10,closed,total = 14.846300 + 0.791795 = 15.638094 kA", compact=True)
    add_equation(doc, "ip,10,max = √2×1.8×15.638094 = 39.808089 kA", compact=True)
    add_equation(doc, "I10,T10-parallel,total = 28.472322 kA；ip = 72.478698 kA（禁止方式）", compact=True)
    add_result_analysis(
        doc,
        "10kV条件职责15.638kA低于31.5kA开关柜等级，因此课程方案不设置限流电抗器。"
        "两台T10并列会把综合值提高到28.472kA并显著抬高峰值，必须通过母联和进线联锁阻止，不能作为可接受的正常运行方式。",
    )

    add_heading(doc, "3.5 全部计算点汇总", 2)
    status_map = {
        "permitted_separate_operation": "允许分列",
        "conditional_pending_system_parallel_permission": "条件并列",
        "conditional_dispatch_permission_only": "条件并列",
        "not_permitted_non_normal_sensitivity": "禁止敏感性",
    }
    point_label_map = {
        "SC-220-BUS-CLOSED": "220kV分段闭合",
        "SC-220-I-SEPARATE": "220kV-I分列",
        "SC-220-II-SEPARATE": "220kV-II分列",
        "SC-35-I-220-CLOSED": "35kV-I上游闭合",
        "SC-35-II-220-CLOSED": "35kV-II上游闭合",
        "SC-35-I-220-SEPARATE": "35kV-I分列",
        "SC-35-II-220-SEPARATE": "35kV-II分列",
        "SC-35-BOTH-TRANSFORMERS-SENSITIVITY": "35kV主变并列敏感",
        "SC-10-I-220-SEPARATE": "10kV-I分列",
        "SC-10-I-220-CLOSED": "10kV-I上游闭合",
        "SC-10-II-220-SEPARATE": "10kV-II分列",
        "SC-10-II-220-CLOSED": "10kV-II上游闭合",
        "SC-10-BOTH-T10-SENSITIVITY": "10kV T10并列敏感",
    }
    sc_rows = []
    for point in sc["points"].values():
        max_total = point.get("conservative_total_symmetrical_current_range_ka", {}).get("maximum")
        sc_rows.append(
            [
                point_label_map.get(point["point_id"], point["point_id"]),
                str(point["voltage_level"]).removesuffix("kV"),
                fmt(point["equivalent_reactance_pu"], 6),
                fmt(point["grid_symmetrical_current_ka"], 3),
                fmt(max_total, 3),
                fmt(point.get("course_total_peak_sensitivity_ka"), 3),
                status_map.get(point["normal_operation_status"], point["normal_operation_status"]),
            ]
        )
    add_table(
        doc,
        "表3-2 短路计算点汇总",
        ["计算点", "电压/kV", "XΣ/pu", "电网Ik/kA", "综合上限/kA", "课程峰值/kA", "方式"],
        sc_rows,
        [3.0, 1.8, 1.9, 1.9, 1.9, 1.9, 3.1],
        6.7,
        nowrap_body_columns={0, 1, 2, 3, 4, 5},
        margin_x=45,
        margin_y=30,
    )
    add_result_analysis(
        doc,
        "汇总表将允许分列、条件并列和禁止敏感性同时列出，是为了说明设备职责的选取过程。"
        "最终控制值为220kV 7.385kA/18.798kA、35kV 16.291kA/41.470kA和10kV 15.638kA/39.808kA；"
        "禁止方式的25.694kA和28.472kA只用于联锁风险审查。",
    )

    add_heading(doc, "第四章 主要电气设备选择与校验", 1, page_break_before=False)
    add_heading(doc, "4.1 校验条件", 2)
    add_body(
        doc,
        "设备选择先由额定电压确定绝缘等级，再由回路最大持续电流确定额定电流，最后以条件性最大短路电流校验开断、关合、峰值和热稳定。"
        "禁止并列场景保留为联锁风险提示，不自动抬高正常必选职责；若最终联锁无法保证，则必须重新选择设备或调整运行方案。",
    )
    add_equation(doc, "Ue ≥ Un；Ie ≥ Imax", "4-1")
    add_equation(doc, "Ibr ≥ Ik；Imaking ≥ ip；Ipeak withstand ≥ ip", "4-2")
    add_equation(doc, "Ith² tr ≥ Ik² t；t = 1.10 s", "4-3")
    switching_duty_paragraph = add_body(
        doc,
        "开断和峰值预筛采用条件性最大职责：220kV 7.385/18.798kA，35kV 16.291/41.470kA，"
        "10kV 15.638/39.808kA。禁止并列敏感性不抬高正常设备必选职责，但必须落实联锁。",
    )
    switching_duty_paragraph.paragraph_format.keep_together = True
    add_equation(doc, "teq = 1.00 + 0.08 = 1.08 s，向上取1.10 s", "4-4")
    add_result_analysis(
        doc,
        "本章所有“通过”均指额定值等级的课程预筛。精确型号、41℃服务条件、厂家温升、机械荷载、CT负担饱和、MOA TOV/能量和接地保护配合"
        "仍未取得完整输入，因此最终状态保持待复核。",
    )

    add_heading(doc, "4.2 额定电压与持续电流", 2)
    add_body(
        doc,
        "额定电流比较采用“候选额定值−所需职责”表示绝对裕度，并同时关注相对裕度。"
        "对接近额定值的35kV进线/母联柜，即使算术上满足，也必须附加环境温度和厂家温升条件。",
    )
    add_table(
        doc,
        "表4-1 代表性设备持续电流校验",
        ["设备/回路", "所需电流/A", "候选额定/A", "裕度/A", "结论"],
        [
            ["220kV线路断路器/\n隔离开关", "735.559", "3150", "2414.441", "预筛通过"],
            ["220kV主变回路", "495.996", "3150", "2654.004", "预筛通过"],
            ["35kV主变进线/母联", "3117.691", "3150", "32.309", "预筛通过；室内≤40℃与厂家温升待核"],
            ["35kV馈线/T10", "最大546.963", "1250", "703.037", "预筛通过"],
            ["10kV T10进线/母联", "1909.586", "2500", "590.414", "预筛通过；室内≤40℃与厂家温升待核"],
            ["10kV SVG馈线", "692.820", "1250", "557.180", "预筛通过"],
        ],
        [4.5, 2.8, 2.7, 2.7, 2.8],
        8.8,
    )
    add_equation(doc, "ΔI35 = 3150 − 3117.691 = 32.309 A", compact=True)
    add_equation(doc, "δ35 = 32.309/3117.691×100% = 1.04%", compact=True)
    add_equation(doc, "ΔI10 = 2500 − 1909.586 = 590.414 A", compact=True)
    add_result_analysis(
        doc,
        "220kV设备额定电流裕度很大，35kV馈线和10kV SVG馈线采用1250A也有明显裕度。"
        "真正控制持续电流的是35kV主变进线/母联3150A等级，只有在开关室受控不高于40℃、通风空调可靠且厂家温升资料确认时才可维持该选择。",
    )

    add_heading(doc, "4.3 开断、关合与动稳定", 2)
    add_body(
        doc,
        "断路器额定开断电流应不小于故障初始对称电流，额定关合电流和设备峰值耐受应不小于课程峰值。"
        "本项目候选等级为220kV 50kA/125kA、35kV和10kV 31.5kA/80kA。",
    )
    add_table(
        doc,
        "表4-2 开断和峰值校验",
        ["电压级", "所需开断/kA", "候选开断/kA", "所需峰值/kA", "候选峰值/kA", "结论"],
        [
            ["220kV", "7.385", "50", "18.798", "125", "预筛通过"],
            ["35kV", "16.291", "31.5", "41.470", "80", "预筛通过"],
            ["10kV", "15.638", "31.5", "39.808", "80", "预筛通过"],
        ],
        [2.1, 2.5, 2.5, 2.5, 2.6, 3.3],
        9.0,
    )
    add_equation(doc, "220kV：50/7.385 = 6.77；125/18.798 = 6.65", compact=True)
    add_equation(doc, "35kV：31.5/16.291 = 1.93；80/41.470 = 1.93", compact=True)
    add_equation(doc, "10kV：31.5/15.638 = 2.01；80/39.808 = 2.01", compact=True)
    add_result_analysis(
        doc,
        "三个电压等级均满足开断和峰值比较。35/10kV裕度约为2倍，足以覆盖当前课程条件职责；"
        "若最终电网等值、变流器模型或允许并列方式使短路水平接近31.5kA，应重新审查限流措施和开关柜等级。",
    )

    add_heading(doc, "4.4 热稳定", 2)
    thermal_method_paragraph = add_body(
        doc,
        "热稳定采用故障电流平方与持续时间的乘积比较。所需热效应按条件性最大短路电流和1.10s等值时间计算；"
        "候选设备能力按额定短时耐受电流及其标称持续时间计算。",
    )
    thermal_method_paragraph.paragraph_format.keep_together = True
    add_equation(doc, "Qk = Ik² teq；Qr = Ith² tr；要求Qr ≥ Qk", "4-5")
    add_equation(doc, "220kV：7.384572²×1.10 = 59.985 kA²·s；50²×3 = 7500 kA²·s", compact=True)
    add_equation(doc, "35kV：16.290983²×1.10 = 291.936 kA²·s；31.5²×4 = 3969 kA²·s", compact=True)
    add_equation(doc, "10kV：15.638094²×1.10 = 269.005 kA²·s；31.5²×4 = 3969 kA²·s", compact=True)
    thermal_rows = []
    for level, ik, ith, tr in (
        ("220kV", 7.384571960617697, 50.0, 3.0),
        ("35kV", 16.290983445518524, 31.5, 4.0),
        ("10kV", 15.638094419605405, 31.5, 4.0),
    ):
        required = ik * ik * 1.10
        available = ith * ith * tr
        thermal_rows.append(
            [level, fmt(ik, 3), fmt(required, 3), f"{fmt(ith,1)}kA/{fmt(tr,0)}s", fmt(available, 1), "预筛通过"]
        )
    add_table(
        doc,
        "表4-3 热稳定I²t校验",
        ["电压级", "Ik / kA", "Ik²t / kA²·s", "候选耐受", "Ith²tr / kA²·s", "结论"],
        thermal_rows,
        [2.3, 2.1, 2.4, 2.7, 2.8, 3.2],
        8.4,
        nowrap_body_columns={0, 1, 2, 3, 4},
        margin_x=75,
        margin_y=65,
    )
    add_result_analysis(
        doc,
        "候选设备额定I²t均远高于课程所需值，因此热稳定不是当前开关设备等级的控制因素。"
        "但该结论依赖1.10s等值时间；最终保护整定或断路器开断时间增加时，应按新的清除时间重新计算。",
    )

    add_heading(doc, "4.5 CT/PT、母线、MOA和接地开关课程预筛", 2)
    add_heading(doc, "4.5.1 母线与导体", 3)
    add_body(
        doc,
        "母线首先按持续载流量选择截面，再按铝导体热稳定常数C=87和1.10s等值时间校验。"
        "220kV管母另利用已冻结的相距、支撑跨距和管形截面进行简化弯曲应力及电晕起始电压计算；"
        "35/10kV矩形母线因缺少支撑机械参数，只进行载流和热稳定。",
    )
    add_equation(doc, "Ith,allow = C S/(1000√t)；Smin = 1000Ik√t/C", "4-6")
    add_equation(
        doc,
        f"S220 = π(100²−90²)/4 = {busbars['BUS-220-TUBE-100-90']['calculated']['cross_section_area_mm2']:.3f} mm²",
        compact=True,
    )
    add_equation(
        doc,
        f"Ith,220 = 87×{busbars['BUS-220-TUBE-100-90']['calculated']['cross_section_area_mm2']:.3f}/(1000√1.10) = "
        f"{busbars['BUS-220-TUBE-100-90']['calculated']['thermal_allowable_current_ka']:.3f} kA",
        compact=True,
    )
    add_equation(doc, "S35 = 3×125×10 = 3750 mm²；Ith,35 = 311.067 kA", compact=True)
    add_equation(doc, "S10 = 2×125×10 = 2500 mm²；Ith,10 = 207.378 kA", compact=True)
    add_equation(
        doc,
        "KT = √[(70−41)/(70−25)] = 0.802773；Iallow,220 = 2000×KT = 1605.546 A",
        "4-7",
        compact=True,
    )
    add_table(
        doc,
        "表4-4 母线课程预筛",
        ["电压级", "课程选择", "持续电流校验", "补充校验"],
        [
            ["220kV", "铝合金管母Φ100/90mm", "1605.546＞735.559A", "热稳123.785＞7.385kA；应力10.048＜70MPa；电晕330.444＞145.492kV"],
            ["35kV", "每相3×125×10mm矩形铝母线", "4194＞3117.691A", "热稳311.067＞16.291kA；机械/电动力待核"],
            ["10kV", "每相2×125×10mm矩形铝母线", "3282＞1909.586A", "热稳207.378＞15.638kA；机械/电动力待核"],
        ],
        [2.2, 4.2, 3.8, 5.3],
        8.2,
        "220kV管母已做简化弯曲和电晕；35/10kV仅完成载流与热稳。支撑共振、连接金具、厂家温升及两级矩形母线机械校验仍待精确资料。",
        keep_note_with_table=True,
    )
    add_body(
        doc,
        f"220kV管母简化电动力按峰值{busbars['BUS-220-TUBE-100-90']['requirements']['required_peak_current_ka']:.3f}kA、"
        f"计算相距{busbars['BUS-220-TUBE-100-90']['calculated']['dynamic']['calculation_phase_spacing_m']:.1f}m和"
        f"支撑跨距8.0m计算。单位长度电磁力为{busbars['BUS-220-TUBE-100-90']['calculated']['dynamic']['base_electromagnetic_force_n_per_m']:.3f}N/m，"
        f"乘1.8动载系数后为{busbars['BUS-220-TUBE-100-90']['calculated']['dynamic']['design_force_n_per_m']:.3f}N/m。",
    )
    add_equation(doc, "f = 2×10⁻⁷ ip²/a；Mmax = 1.8fL²/8；σ = Mmax/W", compact=True)
    add_equation(
        doc,
        f"σ = {busbars['BUS-220-TUBE-100-90']['calculated']['dynamic']['calculated_bending_stress_mpa']:.3f} MPa < 70 MPa",
        compact=True,
    )
    add_equation(
        doc,
        f"Ucorona = {busbars['BUS-220-TUBE-100-90']['calculated']['corona']['critical_disruptive_phase_voltage_kv']:.3f} kV > "
        f"Um/√3 = {busbars['BUS-220-TUBE-100-90']['calculated']['corona']['highest_operating_phase_voltage_kv']:.3f} kV",
        compact=True,
    )
    add_result_analysis(
        doc,
        "220kV管母在持续、热稳定、简化弯曲和电晕四项课程校验中均有明显裕度。"
        "35/10kV矩形母线的热稳定允许值虽远大于短路职责，但没有支撑跨距、固有频率和连接受力就不能宣称机械动稳定已完成。",
    )

    add_heading(doc, "4.5.2 金属氧化物避雷器", 3)
    add_body(
        doc,
        "MOA按持续运行电压和雷电保护裕度预筛。持续运行电压要求与系统最高电压及接地方式有关；"
        "保护裕度采用设备LIWV与避雷器雷电残压之比，课程下限取1.15。",
    )
    add_equation(doc, "Uc ≥ Uc,required；Kp = LIWV/Ures ≥ 1.15", "4-8")
    add_table(
        doc,
        "表4-5 MOA参数与绝缘裕度",
        ["型号", "Ur/Uc/In/残压", "课程校验", "边界"],
        [
            ["YH10W-204/532", "204/159kV；10kA；≤532kV", "Uc余13.508kV；950/532=1.786", "TOV/能量/厂家型号待核"],
            ["YH5WZ-51/134", "51/40.8kV；5kA；≤134kV", "Uc仅余0.300kV（约0.74%）；185/134=1.381", "接地方式、10s TOV、能量、引线压降必核"],
            ["YH5WZ-17/45", "17/13.6kV；5kA；≤45kV", "Uc余1.600kV；75/45=1.667", "TOV/能量/厂家型号待核"],
        ],
        [3.8, 3.8, 4.4, 3.5],
        7.9,
        "LIWV/残压比课程下限取1.15。",
        nowrap_body_columns={0},
        margin_x=75,
        margin_y=65,
    )
    add_equation(doc, "220kV：Uc,required=252/√3=145.492kV；159−145.492=13.508kV；950/532=1.786", compact=True)
    add_equation(doc, "35kV：Uc,required=40.5kV；40.8−40.5=0.300kV；185/134=1.381", compact=True)
    add_equation(doc, "10kV：Uc,required=12kV；13.6−12=1.600kV；75/45=1.667", compact=True)
    add_result_analysis(
        doc,
        "三种MOA均通过课程持续电压和LIWV/残压比检查。35kV Uc只余0.300kV、约0.74%，是最薄弱项；"
        "最终必须结合非有效接地系统10s TOV、能量吸收、线路放电等级和引线压降重新确认，不能仅凭型号字符串定型。",
    )

    add_heading(doc, "4.5.3 CT、PT/CVT与零序CT", 3)
    add_body(
        doc,
        "CT一次变比按持续电流职责向上选择，并同时检查短时电流、I²t和峰值耐受；芯级则按计量、测量、过流、差动及母差功能配置。"
        "PT/CVT先确定一次/二次电压和准确级，容量、铁磁谐振与暂态响应仍由二次专业复核。电缆回路的ZCT承担剩余电流保护，与三相CT职责分开。",
    )
    add_table(
        doc,
        "表4-6 CT/PT课程目标",
        ["对象", "一次变比/电压", "课程目标", "工程待核"],
        [
            ["220kV线路/分段CT", "1000/1A", "0.2S/0.5/5P30，按需PX；50kA/3s、125kA", "负担、饱和、拐点、暂态级"],
            ["220kV主变CT", "600/1A", "0.2S/0.5/5P30/PX；50kA/3s、125kA", "负担、差动配合和精确型号"],
            ["主变中性点CT", "600/1A", "PX限制接地故障+5P30中性点过流", "单相接地/零序职责；不沿用三相短路"],
            ["35kV进线/母联CT", "4000/1A", "0.2S/0.5/5P30/PX；31.5kA/4s、80kA", "负担、饱和和精确型号"],
            ["35kV馈线/T10 CT", "800/1A", "0.2S/0.5/5P30；31.5kA/4s、80kA", "零序与保护配合"],
            ["10kV进线/母联CT", "2500/1A", "0.2S/0.5/5P30/PX；31.5kA/4s、80kA", "负担、饱和和精确型号"],
            ["10kV馈线/SVG CT", "1000/1A", "0.2S/0.5/5P30；31.5kA/4s、80kA", "零序与保护配合"],
            ["35kV电缆馈线ZCT", "100/1A", "5回；一次剩余电流目标400A；5P20或厂家等效", "窗口、屏蔽接地回路、负担和拐点"],
            ["10kV电缆/SVG ZCT", "50/1A", "2回SVG+3类辅助馈线；一次剩余电流目标200A", "窗口、屏蔽接地回路、负担和拐点"],
            ["220/35/10kV PT", "220/√3、35/√3、10/√3kV", "100/√3V测量/保护+100/3V开口三角；0.2或0.5、3P", "容量、铁磁谐振、同期抽取"],
        ],
        [3.2, 3.3, 5.1, 3.9],
        7.7,
    )
    add_body(
        doc,
        "例如35kV进线CT取4000/1A覆盖3117.691A，10kV进线CT取2500/1A覆盖1909.586A，"
        "35kV馈线CT取800/1A覆盖最大546.963A。主变中性点CT不能沿用主变相回路三相短路职责，"
        "其持续、短时、峰值和PX参数须由单相接地及零序专题确定。",
    )
    add_body(
        doc,
        "35kV五回电缆馈线采用100/1A ZCT，10kV SVG及辅助电缆馈线采用50/1A ZCT。"
        "三相电缆芯线必须同穿零序CT窗口，屏蔽层接地回流线不得错误穿入；窗口尺寸、负担、拐点和保护整定在厂家协调阶段闭合。",
    )
    add_result_analysis(
        doc,
        "当前CT/PT/ZCT表完成了课程一次目标和短路等级接口，但尚不是订货清册。"
        "在缺少二次负荷、电缆长度、饱和和铁磁谐振数据时，保留“目标配置、最终型号待核”比给出伪精确型号更符合初步设计深度。",
    )

    add_heading(doc, "4.5.4 接地设备、绝缘与适用性", 3)
    add_body(
        doc,
        "35/10kV三角形绕组或不具中性点的系统通过ZN接地变形成中性点，并串联低电阻把单相接地电流限制在400A和200A。"
        "电阻按相电压除以目标接地电流计算，接地变最小容量按10s等效功率和课程短时过载系数10折算。",
    )
    add_equation(doc, "Uφ=Un/√3；R=Uφ/Ig；Seq=UφIg；ST,min=Seq/Koverload", "4-9")
    add_equation(doc, "35kV：Uφ=20.2073kV；R=50.518Ω；Seq=8.0829MVA；ST,min=808.290kVA", compact=True)
    add_equation(doc, "10kV：Uφ=5.7735kV；R=28.868Ω；Seq=1.1547MVA；ST,min=115.470kVA", compact=True)
    add_body(
        doc,
        "35kV选1000kVA、50.5Ω、400A/10s设备包，10kV选200kVA、28.9Ω、200A/10s设备包。"
        "正常每段各1套；母联合闸前退出受电或故障段接地源，仅保留健康侧1套，重新分列后恢复每段1套。",
    )
    add_table(
        doc,
        "表4-7 绝缘、接地和适用性结论",
        ["项目", "课程采用", "结论或边界"],
        [
            ["绝缘子/设备套管", "252/950、40.5/185、12/75kV；套管630/4000/2500A及本级短时/峰值", "电压、电流、短时和峰值课程检查通过；爬距、机械负荷和精确型号待厂家"],
            ["35kV中性点", "2×1000kVA ZN接地变+低电阻；400A/10s、R≈50.5Ω", "1250A柜；600/1A相CT+400/1A中性点CT；厂家成套待核"],
            ["10kV中性点", "2×200kVA ZN接地变+低电阻；200A/10s、R≈28.9Ω", "1250A柜；300/1A相CT+200/1A中性点CT；厂家成套待核"],
            ["接地源联锁", "母联合闸前退出受电/故障段接地源，仅保留健康侧1套", "禁止母联闭合且两套并联；分列后恢复每段1套"],
            ["接地开关", "220kV线路侧常开ES；35/10kV柜内常开ES", "短时等级随本级设备；感应电流开合和联锁图待核"],
            ["限流电抗器", "不设置", "35/10kV职责低于31.5kA；最终短路超过设备能力时重审"],
            ["高压熔断器", "非主回路设备", "仅PT柜按厂家方案配置一次熔断/二次保护"],
            ["阻波器", "当前不配置", "仅确定电力线载波通信时配置"],
        ],
        [3.4, 5.0, 7.1],
        8.1,
    )
    add_result_analysis(
        doc,
        "接地设备包的电阻和容量算术已闭环，但零序阻抗、温升、NGR结构、CT饱和、保护定值和母联硬/软件联锁仍需厂家统一协调。"
        "绝缘子、套管和接地开关已完成最高电压、LIWV、持续、短时和峰值课程检查；爬距、机械荷载、抗震和感应电流开合等级仍待精确资料。",
    )

    add_heading(doc, "4.6 最终工程复核边界", 2)
    add_body(
        doc,
        "课程设计已经把可由现有输入确定的额定值、数量级和运行逻辑闭合，但初步设计不能替代设备订货和施工专题。"
        "下表逐项说明当前完成深度及关闭最终状态所需的真实输入，避免把“计算脚本有结果”误写成“工程校验全部完成”。",
    )
    add_table(
        doc,
        "表4-8 尚未具备最终校验输入的项目",
        ["项目", "本课程已完成", "最终仍需输入"],
        [
            ["CT", "一次变比、芯级与短路等级目标", "负担、饱和、拐点、暂态级、二次负荷与电缆"],
            ["PT/CVT", "一次/二次电压、准确级和开口三角", "容量、铁磁谐振、同期抽取和精确型号"],
            ["母线/导体", "220kV载流/热稳/简化弯曲/电晕；35/10kV载流/热稳", "35/10kV电动力与支撑机械、厂家温升、共振、无线电干扰和金具"],
            ["绝缘子/套管", "最高电压、LIWV、套管持续/短时/峰值课程检查", "机械破坏负荷、爬距、端子荷载、厂家图"],
            ["MOA", "Ur、Uc、In、残压和LIWV裕度", "接地方式、10s TOV、能量、引线压降和精确型号"],
            ["接地", "接地变+低电阻设备包、馈线CT目标、母联切换逻辑及接地接口", "单相接地/零序、厂家联锁、土壤电阻率、接地网、接触/跨步电压"],
            ["设备环境", "室外41℃；35/10kV室内柜受控≤40℃；≤1000m、d级", "精确型号服务条件和修正系数"],
        ],
        [3.0, 5.0, 7.5],
        8.8,
    )
    add_note_box(
        doc,
        "最终状态",
        "上述额定值校验为课程等级预筛，不代表已取得精确订货型号。报告中应使用“预筛通过、最终型号待厂家资料”"
        "的表述，避免把缺少的数据按0处理或自动判为通过。",
    )
    add_result_analysis(
        doc,
        "优先级最高的后续复核包括35kV 3150A柜温升、LGJ-400/50厂家热额定、35kV MOA TOV/能量、"
        "35/10kV母线机械动稳定以及接地变/NGR与保护联锁。上述项目可能改变精确型号或运行限制，"
        "但在当前课程假设下不改变主接线基本拓扑。",
    )

    add_heading(doc, "4.7 可追溯计算文件", 2)
    add_body(
        doc,
        "正文中的关键数值均可追溯到结构化JSON/CSV结果及相应Python计算模块。"
        "修改YAML输入后，应先重新生成负荷、短路和设备结果，再生成本计算书；不得只在Word表格中手工修改结果。",
    )
    add_table(
        doc,
        "表4-9 计算结果与脚本路径",
        ["内容", "结果文件", "计算脚本"],
        [
            ["负荷、变压器、持续电流", "results/load_and_transformer_\nresults.json", "load_and_transformers/\ncalculate.py"],
            ["回路电流", "results/circuit_currents.csv", "load_and_transformers/\ncalculate.py"],
            ["所用电", "results/station_service_loads.csv", "load_and_transformers/\ncalculate.py"],
            ["短路计算", "results/short_circuit/\nshort_circuit_results.json", "short_circuit/calculate.py"],
            ["设备预筛", "results/equipment_selection/\nequipment_selection_results.json", "equipment_selection/\ncalculate.py"],
        ],
        [3.2, 6.6, 5.7],
        8.1,
        "表中路径均相对于 calculations/ 目录。",
        margin_x=70,
        margin_y=45,
    )
    add_result_analysis(
        doc,
        "可追溯文件用于复算和一致性检查，正式报告仍保留了公式、代入和工程解释，因此即使不运行代码，"
        "读者也能够检查主要计算链；脚本的作用是减少重复运算和数值漂移，而不是用程序输出替代课程设计论证。",
    )
    add_references(doc, REFERENCES, page_break_before=False)

    path = REPORT_DIR / "02_220kV新能源汇集变电所_技术设计计算书.docx"
    finalize_document(doc, path)
    return path


def build_course_summary(data: dict) -> Path:
    doc = new_document("课程设计总结")
    main_sec = start_main_section(doc, SHORT_TITLE + "—课程设计总结")
    add_heading(doc, "课程设计总结", 1, page_break_before=False)
    add_heading(doc, "1 设计任务回顾", 2)
    add_body(
        doc,
        "本次课程设计围绕220kV新能源汇集变电所电气一次部分展开，任务包括主接线、主变和所用变选择、"
        "负荷与持续电流计算、短路计算、主要设备校验、配电装置布置以及三张必交图加一张增强详图。设计过程从任务书"
        "原始数据出发，先冻结运行方式和计算口径，再把同一组基线同步到计算、设备表和CAD图纸。",
    )
    add_heading(doc, "2 主要工作与结果", 2)
    add_bullets(
        doc,
        [
            "确定220/35/10/0.4kV分层接线：前三个电压级均分段分列，0.4kV采用单母线分段暗备用。",
            "完成35kV新能源负荷计算，得到主变校核容量280.286MVA，选2×180MVA，正常负载率77.857%。",
            "明确主变N-1需限发35.780%，同时校核220kV LGJ-400/50导线的41℃课程载流量。",
            "补充任务书未给的10kV方案，选2×31.5MVA T10和2×±12Mvar SVG，补偿后功率因数约0.98147。",
            "采用100MVA标幺法完成220/35/10kV课程短路初算，并区分允许、条件允许和禁止并列场景。",
            "补齐MOA、零序CT、接地开关和35/10kV接地设备包，并固化母联转供时仅保留健康侧1套接地源的联锁。",
            "完成220kV主接线、户外AIS平面、典型断面和L1线路间隔增强详图的DWG/DXF/PDF/PNG底图。",
        ],
    )
    add_heading(doc, "3 设计方法与体会", 2)
    method_paragraph = add_body(
        doc,
        "课程设计中最重要的不是孤立得到一个数值，而是保持任务书、运行方式、计算点、设备职责和图纸表达一致。"
        "例如，35kV和10kV母联的正常状态直接决定短路等值；母联合闸前还必须退出受电段接地源，避免两套"
        "低电阻并联改变接地故障电流；主变N-1限发又会改变220kV线路的持续电流。"
        "如果先画图后冻结这些逻辑，后续会出现大面积返工。因此本项目采用“输入台账—设计基线—计算结果—"
        "设备职责—CAD图纸—文档”的可追溯链条。",
    )
    method_paragraph.paragraph_format.keep_together = True
    boundary_paragraph = add_body(
        doc,
        "另一个重要认识是区分课程假设与工程定值。海拔、污秽、SVG容量、保护时间、变流器短路贡献和AIS几何"
        "都可以为了完成课程设计提出合理且保守的值，但必须写明证据等级和替换条件。对精确设备型号、CT/PT/ZCT"
        "二次参数、35/10kV母线机械、接地网、避雷器和厂家41℃数据，不能因时间有限而虚构“已校验合格”。",
    )
    boundary_paragraph.paragraph_format.keep_together = True
    add_heading(doc, "4 质量控制", 2)
    add_table(
        doc,
        "表1 设计质量控制结果",
        ["检查项", "结果"],
        [
            ["自动化测试", "当前全量70项通过；正式再生成后仍以最新全量结果为准"],
            ["项目健康检查", "当前68条检查通过；正式再生成后重新执行"],
            ["图纸一致性", "四图回路、设备链、编号、净距和几何基线一致"],
            ["公开仓库边界", "任务书、模板、课程资料和个人字段未公开"],
            ["可再生性", "负荷、短路、设备预筛和图纸均有脚本或结构化源数据"],
        ],
        [6.5, 9.0],
        9.8,
    )
    add_heading(doc, "5 后续完善方向", 2, page_break_before=True)
    add_bullets(
        doc,
        [
            "按教师最终口径确认说明书与计算书是否分册手写，以及是否使用右侧教师批阅栏。",
            "依据现行厂家样本复核41℃额定电流、设备外形、基础和端子图。",
            "按GB/T 15544.1-2023补充完整短路计算，并引入新能源、SVG和电动机的精确模型。",
            "补做CT/PT/ZCT负担、避雷器TOV/能量、接地网与接地源联锁、35/10kV母线机械及套管爬距/机械等专项。",
            "根据本项目CAD底图完成最终手绘临摹、打印和提交包核验。",
        ],
        keep_together=True,
    )
    add_note_box(doc, "提交提示", "本总结为公开电子底稿，个人信息和签字在最终手写或私有提交版中补齐。")
    path = REPORT_DIR / "03_220kV新能源汇集变电所_课程设计总结.docx"
    finalize_document(doc, path)
    return path


def add_qa(doc: Document, number: int, question: str, answer: list[str]) -> None:
    add_heading(doc, f"{number}. {question}", 2)
    add_bullets(doc, answer, keep_together=True)


def build_defense_questions(data: dict) -> Path:
    doc = new_document("答辩提纲与问题清单")
    main_sec = start_main_section(doc, SHORT_TITLE + "—答辩提纲")
    add_heading(doc, "答辩提纲与常见问题", 1, page_break_before=False)

    add_heading(doc, "一、90秒开场陈述", 2)
    add_body(
        doc,
        "本设计为220kV新能源汇集变电所电气一次部分初步设计。本站以12回35kV集电线路汇集270MW风电、"
        "光伏和储能，经2×180MVA主变升压后由两回220kV线路送出。主接线采用220kV、35kV、10kV单母线"
        "分段；220kV分段断路器、35/10kV母联正常断开，0.4kV母联闭合并由一台200kVA所用变供电。"
        "35kV综合计算容量为280.286MVA，主变正常负载率77.857%，N-1时需限发35.780%。"
        "为满足无功和辅助系统需求，设置2×31.5MVA T10和2×±12Mvar SVG。短路初算采用100MVA标幺法，"
        "10kV条件性最大综合电流15.638kA，设备按31.5kA等级预筛。35/10kV每段配置ZN接地变+低电阻，"
        "母联转供时仅保留健康侧1套接地源。220kV配电装置采用户外AIS，并完成"
        "三张必交图和SEC-220-L1-01线路间隔增强详图，共四张课程设计底图。",
    )

    add_heading(doc, "二、汇报主线", 2)
    add_bullets(
        doc,
        [
            "题目与原始资料：220/35/10kV、270MW新能源、两回220kV送出。",
            "主接线：为什么选择分段分列，以及各母联正常状态。",
            "容量：2×180MVA主变、2×31.5MVA T10、2×±12Mvar SVG、2×200kVA所用变。",
            "关键计算：280.286MVA、77.857%、35.780%、15.638kA、475.242A。",
            "设备和配电装置：额定等级预筛、CT/PT/ZCT、MOA、母线/绝缘、接地设备包及不设限流电抗器等结论。",
            "边界：哪些是任务书值、哪些是课程假设、哪些必须由厂家和现行标准覆盖。",
        ],
    )

    add_heading(doc, "三、常见答辩问题", 1)
    questions = [
        (
            "为什么220kV侧选单母线分段，而不是双母线？",
            [
                "本期只有两回线路、两台主变和一回预留，单母线分段能把母线故障限制在一段，接线清晰、经济。",
                "双母线检修灵活性更高，但隔离开关多、占地和投资大、倒闸复杂，与本期规模不匹配。",
                "若教师要求母线检修时线路和主变不停电，可升级为双母线不带旁路。",
            ],
        ),
        (
            "2×180MVA主变是否合理？",
            [
                "35kV综合计算容量为280.286MVA，两台均分时每台140.143MVA，负载率77.857%，正常方式合理。",
                "N-1时单台只能覆盖64.220%，但任务书允许新能源汇集站限发，因此通过35.780%限发满足运行。",
            ],
        ),
        (
            "主变N-1和220kV线路N-1谁控制限发？",
            [
                "单回线路在41℃课程载流量475.242A约束下需限发35.390%。",
                "主变N-1需要限发35.780%，略严格；限发后线路电流472.377A，小于课程允许值，余量2.864A。",
                "所以当前控制元件是主变，但线路裕度很小，厂家热额定值必须复核。",
            ],
        ),
        (
            "为什么35kV和10kV母联正常断开？",
            [
                "两台主变和两台T10分别带一段，分列可以限制短路电流并缩小故障影响范围。",
                "母联只在一侧电源隔离且故障清除后用于事故转供，不能把两个健康电源长期并列。",
            ],
        ),
        (
            "35kV或10kV母联事故转供时，接地源怎样切换？",
            [
                "正常分列时每个独立带电母线段各投入1套接地变+低电阻。",
                "母联合闸前先退出受电侧或故障侧接地源，只保留健康电源侧1套；联锁禁止母联闭合且两套同时投入。",
                "母联断开、两段恢复独立供电后，再恢复每段各1套，否则两套低电阻并联会改变目标接地故障电流和保护配合。",
            ],
        ),
        (
            "为什么0.4kV母联正常闭合？",
            [
                "所用变采用暗备用，正常只运行一台，闭合母联后该台可同时带两段低压母线。",
                "另一台所用变及其进线保持断开；切换时先隔离故障进线，再投入备用进线，防止并列。",
            ],
        ),
        (
            "为什么另设35/10.5kV T10，而不用三绕组主变？",
            [
                "任务书给出了220/35kV双绕组和35/10.5kV变压器阻抗，独立T10与给定资料一致。",
                "独立T10可限制10kV短路电流，并使10kV故障不直接迫使主变退出。",
                "2×31.5MVA还能在N-1下承担两套SVG最严吸收工况，负载率约79.17%。",
            ],
        ),
        (
            "SVG容量为什么选2×±12Mvar？",
            [
                "目标功率因数0.98，保守计入10kV辅助无功后需21.8875Mvar。",
                "选24Mvar后容量裕度9.65%，补偿后功率因数约0.98147。",
                "该值是课程冻结值，最终仍由无功电压和谐波专题确定。",
            ],
        ),
        (
            "短路点为什么这样选择？",
            [
                "设备选择主要受各级母线故障控制，因此至少计算220kV、35kV I/II段和10kV I/II段。",
                "另外计算220kV分段断路器条件闭合、健康主变低压侧并列和健康T10并列，用于最大方式和误操作敏感性。",
            ],
        ),
        (
            "为什么禁止并列敏感性不作为设备必选职责？",
            [
                "这些方式在运行规程和联锁中被禁止，不属于正常或允许事故转供方式。",
                "把禁止方式直接作为设备必选职责会掩盖运行边界并导致不必要抬高等级，但其结果必须用于联锁和风险提示。",
            ],
        ),
        (
            "10kV最大短路电流是多少？",
            [
                "单台T10供本段、上游220kV分段条件闭合时，计SVG课程上界为15.638kA。",
                "两台健康T10并列的禁止敏感性为28.472kA；正常设备按31.5kA开断等级预筛。",
            ],
        ),
        (
            "设备热稳定如何校验？",
            [
                "采用Ik²t≤Ith²tr，热等值时间取1.10s。",
                "220kV候选50kA/3s，35kV和10kV候选31.5kA/4s，按条件性最大Ik计算均有较大I²t裕度。",
            ],
        ),
        (
            "35kV 3150A和10kV 2500A选型有什么风险？",
            [
                "35kV主变进线需3117.691A，3150A柜仅余32.309A、约1.04%；10kV进线需1909.586A，2500A柜余590.414A。",
                "两级室内开关柜均按受控≤40℃。35kV余量尤其小，必须依赖厂家温升确认；10kV虽有较大余量，精确柜型和散热条件仍需闭合。",
            ],
        ),
        (
            "CT/PT为什么没有写成精确订货型号？",
            [
                "一次安装位置和变比量级可以由持续电流确定，但准确级、芯数、暂态性能和二次负荷依赖保护、计量和电缆清册。",
                "缺少这些输入时写精确型号会造成虚假闭环，因此本阶段只冻结配置原则。",
            ],
        ),
        (
            "任务书没有逐字列避雷器和零序CT，为什么仍要配置？",
            [
                "任务书要求主要电气设备和互感器完整选择，MOA属于过电压保护完整性，零序CT属于电缆单相接地保护接口。",
                "7回35kV架空馈线入口配置MOA；5回35kV电缆和10kV电缆/SVG回路配置零序CT，职责与三相CT分开。",
                "课程阶段只冻结位置和参数目标，MOA的TOV/能量及ZCT窗口、负担、拐点和精确型号仍待专题。",
            ],
        ),
        (
            "35kV和10kV母线是否已经完成动稳定校验？",
            [
                "没有笼统宣称完成。两级矩形母线当前只完成载流和1.10s热稳定课程预筛。",
                "电动力、支撑绝缘子受力、连接金具和厂家温升仍待精确布置与厂家参数；220kV管母另完成了简化弯曲和电晕预筛。",
            ],
        ),
        (
            "为什么220kV配电装置选户外AIS？",
            [
                "本站回路数量有限、场地课程假设允许，AIS经济、直观且便于扩建和手绘表达。",
                "GIS占地小、环境适应性强，但投资高且本题无紧凑场地约束。",
            ],
        ),
        (
            "图纸中的尺寸能否施工使用？",
            [
                "不能。站区145×90m、14m节距、设备轮廓和构架标高均为课程设计几何。",
                "净距来自课程讲义，正式工程仍须按DL/T 5352-2018全文、真实场址和厂家外形基础图复核。",
            ],
        ),
    ]
    for index, (question, answer) in enumerate(questions, 1):
        add_qa(doc, index, question, answer)

    add_heading(doc, "四、答辩时容易说错的边界", 1, page_break_before=False)
    add_table(
        doc,
        "表1 表述边界",
        ["不建议说", "建议说"],
        [
            ["设备已经最终定型", "额定值等级预筛通过，精确厂家型号和41℃条件待核"],
            ["短路电流符合现行标准最终值", "采用任务书标幺X-only课程初算，最终按GB/T 15544.1-2023复核"],
            ["图纸可以施工", "图纸为课程设计和手绘临摹底图"],
            ["35kV/10kV母联可长期闭合", "仅在故障源隔离且故障清除后事故转供"],
            ["母联合闸后两套接地源仍可同时投入", "合母联前退出受电/故障段接地源，仅保留健康侧1套"],
            ["35/10kV母线机械动稳定已全部完成", "当前只完成载流和热稳；电动力与支撑机械待复核"],
            ["CT/PT型号已选定", "一次变比、芯级/准确级目标已冻结；二次负担、暂态性能和最终型号等待保护计量清册"],
        ],
        [7.4, 8.1],
        9.3,
    )
    path = REPORT_DIR / "05_220kV新能源汇集变电所_答辩问题清单.docx"
    finalize_document(doc, path)
    return path


def finalize_document(doc: Document, path: Path) -> None:
    add_update_fields_setting(doc)
    props = doc.core_properties
    props.title = TITLE
    props.subject = "发电厂电气部分课程设计"
    props.author = "课程设计项目组（公开版）"
    props.last_modified_by = "课程设计项目组（公开版）"
    props.comments = ENGINEERING_NOTE
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    data = load_data()
    outputs = [
        build_design_description(data),
        build_calculation_book(data),
        build_course_summary(data),
        build_defense_questions(data),
    ]
    for path in outputs:
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()
